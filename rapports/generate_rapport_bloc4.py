"""Génération du rapport professionnel Bloc 3 — Épreuve E4 (C14, C15, C16, C17, C18, C19)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

OUTPUT_PATH = Path(__file__).parent / "Rapport_Bloc4_Application_Web_IA.docx"
ASSETS_DIR = Path(__file__).parent / "assets_bloc4"


def set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2E74B5")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D9E2F3")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_bullet(doc, title: str, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"• {title} : ")
    run.bold = True
    p.add_run(text)


def add_figure(doc, filename: str, caption: str, width_cm=13):
    doc.add_picture(str(ASSETS_DIR / filename), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        doc.styles[f"Heading {level}"].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # ── PAGE DE GARDE ──
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Rapport professionnel — Développement d'une\n"
        "application intégrant un service d'intelligence artificielle"
    )
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Titre professionnel Développeur en Intelligence Artificielle — Bloc 3, Épreuve E4\n"
        "Bloc4_app — application Django consommant les services de données (Bloc1) "
        "et de classification (Bloc3)"
    )
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Projet support : Classification de Tendance Crypto\n"
        "Certification : RNCP37827 — Développeur en Intelligence Artificielle\n"
        "Compétences visées : C14, C15, C16, C17, C18, C19\n"
        "Dépôt : crypto-certification (GitHub)"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_page_break()

    # ── SOMMAIRE ──
    doc.add_heading("Sommaire", level=1)
    for item in [
        "1. Introduction",
        "2. Analyser le besoin (C14)",
        "3. Concevoir le cadre technique (C15)",
        "4. Conduite de projet (C16)",
        "5. Développer les composants et interfaces (C17)",
        "6. Intégration continue (C18)",
        "7. Livraison continue (C19)",
        "8. Plan de démonstration",
        "9. Bilan et limites assumées",
        "10. Conclusion",
        "11. Sources consultées",
        "Annexe — Repères des sources par thématique",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Les rapports précédents portaient sur la collecte et le stockage des données de marché "
        "(E1, Bloc1_data), la veille sur les services d'IA (E2, Bloc2_veille) et la mise en "
        "service du modèle de classification de tendance (E3, ml-api). Celui-ci porte sur "
        "l'application qui consomme ces deux services : Bloc4_app, un webapp Django hébergeant "
        "trois modules — comptes utilisateurs, tableau de bord et classification à la demande."
    )
    doc.add_paragraph(
        "Point de vigilance à formuler explicitement : Bloc4_app ne définit aucun schéma de "
        "données métier propre — il consomme en lecture les données déjà décrites en E1 "
        "(paires, OHLCV, prédictions stockées) via l'API data-api, et déclenche des "
        "classifications à la demande via l'API ml-api décrite en E3. La seule donnée que "
        "Bloc4_app possède réellement est le compte utilisateur (modèle User natif de Django)."
    )
    doc.add_paragraph(
        "Le besoin fondateur tient en une phrase : « En tant qu'investisseur crypto, je veux "
        "une vision synthétique de la tendance probable du Bitcoin, afin d'orienter mes "
        "décisions à court terme sans avoir à interroger les APIs moi-même. » Bloc4_app y répond "
        "en restituant les prédictions déjà calculées (Dashboard, Charts) et en permettant de "
        "déclencher une classification fraîche à la demande (Forecast)."
    )

    # ══════════════════════════════════════════════════════════════════
    # 2. ANALYSE DU BESOIN — C14
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("2. Analyser le besoin (C14)", level=1)

    doc.add_heading("2.1 User stories", level=2)
    add_bullet(doc, "US1 (fondatrice)",
        "En tant qu'investisseur, je veux connaître la tendance probable du Bitcoin, afin "
        "d'orienter mes décisions à court terme. Couverte par les pages Dashboard et Forecast.")
    add_bullet(doc, "US2",
        "En tant qu'utilisateur, je veux consulter l'historique des prix et des prédictions "
        "passées par paire, afin de juger de la fiabilité du modèle avant de m'y fier. → page "
        "Charts (dashboard/charts.html).")
    add_bullet(doc, "US3",
        "En tant qu'utilisateur, je veux déclencher une classification à la demande sur "
        "l'horizon de mon choix (journalier ou horaire), afin d'obtenir une prédiction fraîche "
        "plutôt que la dernière valeur stockée. → page Forecast (forecast/classify.html).")
    add_bullet(doc, "US4",
        "En tant qu'utilisateur, je veux un accès personnel protégé par mot de passe, afin que "
        "mes consultations ne soient pas ouvertes à n'importe qui. → module accounts.")
    doc.add_paragraph(
        "Ces user stories sont dérivées de l'usage réel du projet (données déjà exposées par "
        "l'API Bloc1, modèle déjà exposé par l'API Bloc3) plutôt qu'issues d'un atelier de "
        "cadrage formel préalable avec des utilisateurs finaux — limite assumée, cohérente avec "
        "un projet de certification individuel."
    )

    doc.add_heading("2.2 Données consommées", level=2)
    doc.add_paragraph(
        "Bloc4_app ne modélise aucune donnée métier en base locale : le Modèle Conceptuel de "
        "Données complet (CURRENCY, TRADING_PAIR, OHLCV_*, PREDICTION_*) appartient à l'API "
        "data-api et a été présenté dans le rapport E1. Bloc4_app en consomme un sous-ensemble "
        "en lecture seule, via HTTP/JWT, résumé ci-dessous :"
    )
    add_styled_table(doc,
        ["Endpoint consommé", "Donnée obtenue", "Utilisé par"],
        [
            ["GET /trading_pairs/trading_pair_by_currency_symbols", "Identifiant de la paire (BTC/USDT, BTC/USD)", "dashboard/services.py::DashboardService"],
            ["GET /ohlcv/{granularity}_by_trading_pair_id", "Historique de prix (open/high/low/close/volume)", "Dashboard, Charts, Monitorage"],
            ["GET /predictions/{granularity}_by_trading_pair_id", "Prédictions déjà stockées (classe, confiance)", "Dashboard, Charts, Monitorage"],
            ["POST /classify/classify_daily ou _hourly (ml-api)", "Classification calculée à la demande", "forecast/services.py::ForecastService"],
        ],
        col_widths=[7, 6, 4],
    )
    doc.add_paragraph(
        "La seule donnée réellement possédée par Bloc4_app est le compte utilisateur — modèle "
        "User natif de Django (accounts/models.py), sans extension métier."
    )

    doc.add_heading("2.3 Parcours utilisateur", level=2)
    add_styled_table(doc,
        ["Étape", "Écran", "Action"],
        [
            ["1", "Connexion / inscription", "S'authentifier (ou créer un compte) — accounts/login, accounts/register"],
            ["2", "Tableau de bord", "Vue d'ensemble : dernier prix et dernière prédiction stockée, par paire — dashboard/index.html"],
            ["3", "Graphiques", "Sélection paire + granularité → historique OHLCV et prédictions superposés — dashboard/charts.html"],
            ["4", "Classification à la demande", "Sélection paire, granularité, appel réel à ml-api → résultat affiché — forecast/classify.html"],
        ],
        col_widths=[2, 5, 10],
    )

    doc.add_heading("2.4 Wireframes (description structurelle)", level=2)
    doc.add_paragraph(
        "Pas de maquette graphique produite en amont — les gabarits ci-dessous décrivent la "
        "structure réelle des templates Django, cohérente avec le tableau des parcours "
        "ci-dessus :"
    )
    add_bullet(doc, "Login / Register", "formulaire centré (username, mot de passe), lien croisé vers l'autre formulaire, messages d'erreur Django au-dessus du formulaire.")
    add_bullet(doc, "Dashboard", "navbar (Dashboard | Charts | Forecast), une carte par paire suivie (prix + label de tendance).")
    add_bullet(doc, "Charts", "deux sélecteurs (paire, granularité), bouton « Charger », graphique Plotly (prix + prédictions).")
    add_bullet(doc, "Forecast", "trois champs (paire, granularité, formulaire ClassifyForm), résultat affiché sous forme de tableau (date, classe, confiance).")

    doc.add_heading("2.5 Accessibilité", level=2)
    doc.add_paragraph(
        "Standard visé : WCAG 2.1 AA. Évaluation honnête plutôt qu'un audit formel — l'application "
        "utilise une structure sémantique HTML5 réelle (balises <nav> et <main> dans "
        "templates/base.html) et hérite des contrastes par défaut de Bootstrap 5, mais l'examen "
        "direct des six templates montre que plusieurs objectifs affichés dans une version "
        "antérieure de la documentation ne sont pas encore tenus :"
    )
    add_styled_table(doc,
        ["Critère WCAG", "État réel constaté"],
        [
            ["1.1.1 Contenu non textuel", "Non tenu : aucun attribut alt ni aria-label dans les templates — le graphique Plotly (charts.html) n'a pas d'alternative textuelle."],
            ["1.3.1 Information et relations", "Tenu : structure sémantique <nav>/<main>, formulaires Django avec labels associés."],
            ["1.4.3 Contraste minimum", "Probablement tenu par défaut Bootstrap 5, mais non mesuré (pas d'outil de contrôle de contraste exécuté)."],
            ["2.4.1 Contourner des blocs", "Non tenu : aucun lien « Aller au contenu » dans templates/base.html."],
            ["3.3.1 Identification des erreurs", "Tenu : erreurs de formulaire Django affichées au champ concerné (LoginForm, RegisterForm)."],
        ],
        col_widths=[5, 11],
    )
    doc.add_paragraph(
        "Deux limites concrètes à corriger avant une diffusion au-delà d'un usage personnel : "
        "l'absence d'attributs alt/aria-label et l'absence de lien d'évitement — reprises dans "
        "le bilan plutôt que présentées comme acquises."
    )
    doc.add_paragraph(
        "Anomalie trouvée en capturant la page de connexion pour ce rapport (cf. Figure 1) : le "
        "champ mot de passe affichait le libellé anglais « Password » alors que le champ nom "
        "d'utilisateur, lui, était bien traduit (« Nom d'utilisateur »). Cause identifiée dans "
        "accounts/forms.py : LoginForm redéclare le champ password sans label explicite, "
        "perdant la traduction que AuthenticationForm applique automatiquement au seul champ "
        "username. Corrigé par l'ajout de label=\"Mot de passe\" — la capture ci-dessous reflète "
        "l'état corrigé."
    )
    add_figure(doc, "crop_login.png",
        "Figure 1 — Page de connexion réelle (après correctif du libellé du mot de passe).")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 3. CONCEPTION TECHNIQUE — C15
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("3. Concevoir le cadre technique (C15)", level=1)

    doc.add_heading("3.1 Architecture et flux de données", level=2)
    doc.add_paragraph(
        "Bloc4_app est un service parmi quatre, orchestrés par docker-compose : "
        "Utilisateur (navigateur, port 8090) → Django (webapp) → JWT → data-api (port 8001, "
        "Bloc1) et ml-api (port 8002, Bloc3). data-api lit/écrit PostgreSQL ; ml-api charge des "
        "modèles .pkl. Bloc4_app n'accède jamais directement à PostgreSQL ni aux modèles — tout "
        "passe par les deux APIs, authentifiées par un compte de service (JWT) distinct de la "
        "session utilisateur (cookie Django)."
    )
    doc.add_paragraph(
        "Trois pages (Dashboard, Charts, Monitorage — cf. rapport E3 pour cette dernière) "
        "interrogent uniquement data-api ; seule la page Forecast interroge ml-api, en plus de "
        "data-api pour l'affichage. Cette séparation reflète le fait que Forecast est la seule "
        "page qui déclenche un calcul, les autres se contentant de restituer des valeurs déjà "
        "stockées."
    )

    doc.add_heading("3.2 Hébergement et éco-conception", level=2)
    doc.add_paragraph(
        "Contrairement à un scénario de mise à disposition publique (PaaS type Render), "
        "Bloc4_app est ici auto-hébergé via docker-compose sur une seule machine — choix "
        "assumé et cohérent avec le périmètre d'un projet de certification, pas encore destiné "
        "à un public externe. Éléments d'éco-conception réels, vérifiés dans le code :"
    )
    add_bullet(doc, "Image Docker légère", "python:3.11-slim, dépendances figées et installées sans le groupe dev en production (uv sync --frozen --no-dev --no-install-project, app.Dockerfile).")
    add_bullet(doc, "Pas de sur-approvisionnement", "un seul conteneur webapp (gunicorn), pas de réplication ni de scaling horizontal pour un usage à trafic faible et prévisible.")
    add_bullet(doc, "Requêtes API bornées", "chaque vue interroge data-api/ml-api pour un nombre fixe de paires (2 à ce stade), pas de pagination ni de sur-fetch.")
    doc.add_paragraph(
        "Argument qualitatif, pas un chiffre mesuré — aucun outil de mesure d'empreinte carbone "
        "par requête n'a été mis en œuvre, cohérent avec la même limite déjà posée sur les "
        "rapports précédents."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 4. CONDUITE DE PROJET — C16
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("4. Conduite de projet (C16)", level=1)
    doc.add_paragraph(
        "Projet individuel, sans équipe formelle — la conduite de projet est adaptée en "
        "conséquence plutôt que de simuler des rituels d'équipe qui n'auraient pas de sens pour "
        "un développeur seul."
    )
    doc.add_heading("4.1 Outils de pilotage réels", level=2)
    add_bullet(doc, "Backlog", "GitHub Issues du dépôt crypto-certification (8 issues ouvertes à date), qui tracent des limites identifiées pendant les rapports E2/E3 — pas un backlog dédié à la construction de Bloc4_app, suivie directement par commit.")
    add_bullet(doc, "Branches", "deux branches de travail réelles (bloc1_data, bloc2-bloc3-mise-en-service), divergentes depuis le commit initial du dépôt ; non encore fusionnées dans main à ce stade — limite reprise dans le bilan plutôt que dissimulée.")
    add_bullet(doc, "Board Kanban", "envisagé (Backlog → In Progress → Review → Done) mais jamais activé dans les faits — le suivi réel s'est fait par les commits et les issues, pas par un board GitHub Projects.")
    doc.add_heading("4.2 Rythme de travail réel", level=2)
    doc.add_paragraph(
        "Pas de cérémonie agile formelle (pas de daily standup, pas de planning hebdomadaire) — "
        "le suivi s'est fait par itérations courtes, chaque commit correspondant à un incrément "
        "testé avant de passer au suivant, la convention de message (feat(bloc4): ..., "
        "fix(bloc4): ..., test(bloc4): ...) servant de journal de bord :"
    )
    add_styled_table(doc,
        ["Commit", "Contenu"],
        [
            ["feat(bloc4): monitorage de dérive du modèle en production (C11)", "dashboard/metrics.py + vue monitoring"],
            ["fix(bloc4): timeout + retry sur token expiré dans ForecastService (C10)", "robustesse de l'intégration ml-api"],
            ["test(bloc4): ajoute les tests de monitoring_view oubliés au commit C11", "complète la couverture après coup"],
        ],
        col_widths=[9, 7],
    )
    doc.add_paragraph(
        "C'est une conduite de projet réelle mais légère, assumée comme telle plutôt que "
        "présentée comme un Scrum ou un Kanban actif qu'elle n'est pas."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 5. RÉALISATION — C17
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("5. Développer les composants et interfaces (C17)", level=1)

    doc.add_heading("5.1 Composants développés", level=2)
    add_styled_table(doc,
        ["Fichier", "Lignes", "Rôle"],
        [
            ["accounts/views.py + forms.py", "45 + 24", "Inscription (UserCreationForm), connexion (AuthenticationForm), déconnexion"],
            ["dashboard/views.py + services.py", "96 + 62", "Vue d'ensemble, graphiques, monitorage de dérive ; DashboardService (appels data-api)"],
            ["dashboard/metrics.py", "88", "Calcul de la dérive prédit/réalisé (cf. rapport E3, C11)"],
            ["forecast/views.py + forms.py + services.py", "33 + 15 + 54", "Formulaire de classification, ForecastService (appels ml-api)"],
            ["crypto_app/settings.py + urls.py", "93 + 9", "Configuration Django, routage racine"],
        ],
        col_widths=[7, 3, 6],
    )
    add_figure(doc, "crop_dashboard.png",
        "Figure 2 — Dashboard réel (dashboard/index.html), paires BTC/USDT et BTC/USD, "
        "prix et dernière prédiction stockée lus en direct depuis data-api.")

    doc.add_heading("5.2 Communication inter-services", level=2)
    doc.add_paragraph(
        "Chaque module de données (DashboardService, ForecastService) s'authentifie séparément "
        "auprès de l'API qu'il consomme, avec un compte de service dédié (identifiants en "
        "variable d'environnement) : POST /authentification/login → jeton JWT mis en cache en "
        "mémoire pour la durée de vie de l'instance du service, puis transmis en en-tête Bearer."
    )
    add_figure(doc, "crop_charts.png",
        "Figure 3 — Page Charts réelle (dashboard/charts.html), BTC/USDT journalier sur 30 "
        "jours — bougies OHLCV et prédictions (losanges rouges) rendus via Plotly, données "
        "lues depuis data-api au clic sur « Charger ».")
    add_figure(doc, "crop_forecast.png",
        "Figure 4 — Page Forecast réelle (forecast/classify.html) : classification à la "
        "demande sur BTC/USDT journalier, appel effectif à ml-api au moment de la capture "
        "(résultat DOWN, confiance 62,3 %, daté du jour même) — preuve que l'intégration "
        "fonctionne de bout en bout, pas seulement en test mocké.")

    doc.add_heading("5.3 Anomalie trouvée et corrigée pendant ce rapport", level=2)
    doc.add_paragraph(
        "En relisant DashboardService avant rédaction — par symétrie avec le correctif déjà "
        "appliqué à ForecastService en E3 (C10) — un même défaut a été trouvé :"
    )
    add_bullet(doc, "Absence de timeout (bug réel)",
        "les quatre appels HTTP de DashboardService (_get_token, get_trading_pair, get_ohlcv, "
        "get_predictions) n'avaient aucun timeout — un data-api qui répond lentement pouvait "
        "bloquer indéfiniment le thread Django. Corrigé par l'ajout d'un timeout explicite "
        "(10s) sur chaque appel, avec deux tests dédiés qui vérifient que le paramètre est bien "
        "transmis (test_get_token_passes_timeout, test_get_ohlcv_passes_timeout)."
    )
    doc.add_paragraph(
        "Point non corrigé, documenté comme limite plutôt que traité en urgence pour ne pas "
        "élargir le périmètre de ce correctif : contrairement à ForecastService, DashboardService "
        "ne capture aucune exception réseau (requests.RequestException) — une panne de data-api "
        "provoquerait une erreur 500 Django non gérée plutôt qu'un message explicite à "
        "l'utilisateur. Repris dans le bilan."
    )

    doc.add_heading("5.4 Gestion des droits d'accès", level=2)
    doc.add_paragraph(
        "Les vues Dashboard, Charts, Monitorage et Forecast sont protégées par le décorateur "
        "@login_required (redirection vers /login/ si non authentifié, vérifié par "
        "tests/test_accounts.py::TestAccessControl). Les mots de passe sont hashés via le "
        "backend PBKDF2 par défaut de Django ; le middleware CSRF est actif sur tous les "
        "formulaires POST."
    )

    doc.add_heading("5.5 Sécurité — revue OWASP Top 10", level=2)
    doc.add_paragraph(
        "Comme pour ml-api (rapport E3), aucun audit tiers formalisé — auto-évaluation contre "
        "les dix catégories, réalisée en lisant effectivement crypto_app/settings.py et les "
        "vues plutôt que par supposition :"
    )
    add_styled_table(doc,
        ["Catégorie OWASP", "Statut sur Bloc4_app"],
        [
            ["A01 — Contrôle d'accès défaillant", "Couvert pour les pages métier (@login_required, vérifié par tests) ; l'interface d'administration Django (/admin/) reste activée sans durcissement particulier — à revoir avant exposition publique."],
            ["A02 — Défaillances cryptographiques", "Partiel : mots de passe hashés (PBKDF2, validateurs Django par défaut) ; secret Django et identifiants inter-services en variables d'environnement ; le chiffrement en transit (TLS) n'est pas géré par Django lui-même, dépend d'un reverse-proxy absent de ce déploiement docker-compose."],
            ["A03 — Injection", "Couvert : aucune requête SQL directe (Bloc4_app n'a pas de modèle métier propre) ; les identifiants de paire proviennent de listes déroulantes bornées (ChoiceField), pas d'un champ texte libre."],
            ["A04 — Conception non sécurisée", "Partiel : pas de limitation de tentatives de connexion (rate-limiting applicatif absent sur /login/), comportement par défaut de Django."],
            ["A05 — Mauvaise configuration de sécurité", "Vérifiée en pratique : DEBUG vaut True par défaut si la variable d'environnement n'est pas positionnée (crypto_app/settings.py) — docker-compose la force à False pour le service webapp, mais ce n'est pas garanti par le code lui-même."],
            ["A06 — Composants vulnérables ou obsolètes", "Dépendances figées et reproductibles (uv.lock), pas de scan automatisé de vulnérabilités en CI — limite déjà identifiée côté ml-api en E3, toujours vraie ici."],
            ["A07 — Identification et authentification défaillantes", "Partiel : validateurs de mot de passe Django actifs (longueur, similarité, mots de passe courants) ; pas de verrouillage de compte après échecs répétés."],
            ["A08 — Manque d'intégrité des données et logiciels", "Partiel : image construite depuis un Dockerfile versionné avec dépendances verrouillées, mais contrairement à ml-api, la chaîne CI ne construit ni ne publie d'image webapp vers un registre (cf. §7)."],
            ["A09 — Carence de journalisation et de surveillance", "Non couvert : contrairement à data-api et ml-api (instrumentées Prometheus, cf. rapport E3), Bloc4_app n'expose aucune métrique — vérifié dans monitoring/prometheus.yml, qui ne liste que data-api et ml-api comme cibles de scrape."],
            ["A10 — Falsification de requête côté serveur (SSRF)", "Faible : les seules requêtes sortantes ciblent des URLs fixes de configuration (API_E1_BASE_URL, API_E3_BASE_URL), jamais une URL fournie par l'utilisateur."],
        ],
        col_widths=[5, 11],
    )
    doc.add_paragraph(
        "Le point le plus net de cette revue est l'absence totale d'instrumentation Prometheus "
        "sur le webapp (A09) — une asymétrie réelle avec data-api/ml-api, à corriger en priorité "
        "si Bloc4_app doit un jour être surveillé en production."
    )

    doc.add_heading("5.6 Dépendances et variables d'environnement", level=2)
    doc.add_paragraph(
        "Dépendances principales (pyproject.toml) : django, gunicorn (serveur WSGI de "
        "production), requests, python-dotenv, psycopg2-binary. Groupe dev séparé (pytest, "
        "pytest-django, pytest-cov, ruff) jamais installé en production (--no-dev dans le "
        "Dockerfile)."
    )
    add_styled_table(doc,
        ["Variable", "Usage"],
        [
            ["DJANGO_SECRET_KEY / DEBUG / ALLOWED_HOSTS", "Configuration cœur Django"],
            ["DJANGO_DB_* (ENGINE, NAME, USER, PASSWORD, HOST, PORT)", "Connexion PostgreSQL (auth_user, sessions)"],
            ["API_E1_BASE_URL / API_E1_SCRIPT_USERNAME / API_E1_SCRIPT_PASSWORD", "Accès en lecture à data-api (Bloc1)"],
            ["API_E3_BASE_URL / API_E3_USERNAME / API_E3_PASSWORD", "Accès à ml-api pour la classification à la demande (Bloc3)"],
        ],
        col_widths=[8, 8],
    )

    doc.add_heading("5.7 Tests", level=2)
    doc.add_paragraph(
        "56 tests (pytest + pytest-django), exécutés avec succès sur SQLite en mémoire "
        "(configuration CI) :"
    )
    add_styled_table(doc,
        ["Fichier", "Tests", "Vérifie"],
        [
            ["test_accounts.py", "14", "Connexion, inscription (doublon, mots de passe non concordants), déconnexion, contrôle d'accès sur les 4 pages protégées"],
            ["test_dashboard.py", "11", "Rendu du dashboard, cas sans prédiction/sans paire, page Charts, API JSON interne, monitorage de dérive (cf. E3)"],
            ["test_forecast.py", "6", "Formulaire de classification, succès, erreur API, erreur réseau, accès protégé"],
            ["test_metrics.py", "10", "Règle de classification réalisée (UP/DOWN/STABLE) et calcul de dérive (cf. E3, C11)"],
            ["test_services.py", "15", "DashboardService et ForecastService : authentification, timeout, retry sur 401, erreurs réseau"],
        ],
        col_widths=[6, 3, 9],
    )
    doc.add_paragraph(
        "Suite exécutée avec les variables d'environnement de la CI (SQLite en mémoire) : "
        "56 tests réussis. Sources versionnées sur le dépôt Git distant."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 6. INTÉGRATION CONTINUE — C18
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("6. Intégration continue (C18)", level=1)
    doc.add_paragraph(
        "Bloc4_app partage la chaîne GitHub Actions du dépôt (.github/workflows/ci.yml) avec "
        "Bloc1_data et Bloc3_ml — un job dédié (test-bloc4) et un job de lint partagé :"
    )
    add_bullet(doc, "Déclencheurs", "push et pull request sur main et develop.")
    add_bullet(doc, "test-bloc4", "installation via uv sync --dev, exécution de pytest tests/ -v --cov, avec les variables SQLite en mémoire et des URLs d'API simulées (les tests mockent les appels HTTP réels) ; upload du rapport de couverture en artefact.")
    add_bullet(doc, "lint", "ruff check Bloc4_app/ --ignore=E501, dans un job commun aux trois blocs.")
    doc.add_paragraph(
        "Contrairement à test-bloc3, le job test-bloc4 n'est suivi d'aucune étape de build ou "
        "de déploiement d'image dans la chaîne actuelle (cf. §7) — asymétrie assumée, reprise "
        "dans le bilan."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 7. LIVRAISON CONTINUE — C19
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("7. Livraison continue (C19)", level=1)
    doc.add_heading("7.1 Packaging", level=2)
    doc.add_paragraph(
        "Image Docker (app.Dockerfile, python:3.11-slim) : dépendances installées via "
        "uv sync --frozen --no-dev --no-install-project (versions figées depuis uv.lock, "
        "pas de dépendances de développement en production), lancement via entrypoint.sh "
        "(migrate + collectstatic + gunicorn)."
    )
    doc.add_heading("7.2 Déploiement — périmètre assumé", level=2)
    doc.add_paragraph(
        "Le service webapp est déclaré dans docker-compose.yml (build depuis app.Dockerfile, "
        "port 8090→8080, dépend de db/data-api/ml-api). Contrairement à ml-api (cf. rapport E3, "
        "jobs build-ml-api/deploy-ml-api qui publient une image sur ghcr.io après succès des "
        "tests), la chaîne CI actuelle ne construit ni ne publie d'image webapp vers un "
        "registre : le déploiement de Bloc4_app se fait uniquement par docker compose up "
        "--build en local, sans vérification automatisée de la construction de l'image en CI. "
        "C'est un écart réel avec la pratique déjà appliquée côté ml-api, repris dans le bilan "
        "comme axe d'harmonisation plutôt que présenté comme équivalent."
    )
    doc.add_heading("7.3 Configuration versionnée", level=2)
    doc.add_paragraph(
        "app.Dockerfile, entrypoint.sh, docker-compose.yml (service webapp) sont tous versionnés "
        "sur le dépôt distant."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 8. PLAN DE DÉMONSTRATION
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("8. Plan de démonstration", level=1)
    doc.add_heading("1. L'application en local", level=2)
    doc.add_paragraph(
        "docker compose up -d — inscription, connexion, Dashboard (prix + prédiction stockée "
        "par paire), Charts (historique OHLCV + prédictions), Forecast (classification à la "
        "demande, appel réel à ml-api)."
    )
    doc.add_heading("2. La chaîne CI en direct", level=2)
    doc.add_paragraph(
        "Un commit trivial poussé sur une branche : montrer le job test-bloc4 (56 tests) et le "
        "job lint se dérouler sur GitHub Actions."
    )
    doc.add_heading("3. Le correctif DashboardService", level=2)
    doc.add_paragraph(
        "Montrer le diff (ajout du timeout), les deux tests associés, et la suite complète "
        "verte — exemple concret de bug trouvé et corrigé en relisant le code avant rédaction "
        "du rapport."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 9. BILAN ET LIMITES ASSUMÉES
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("9. Bilan et limites assumées", level=1)
    add_bullet(doc, "Accessibilité",
        "structure sémantique HTML5 réelle, mais pas d'attributs alt/aria-label ni de lien "
        "d'évitement — deux limites concrètes identifiées, pas d'audit WCAG formel mené.")
    add_bullet(doc, "Conduite de projet",
        "pas de board Kanban actif malgré l'intention initiale ; deux branches de travail non "
        "encore fusionnées dans main — à traiter avant une diffusion plus large.")
    add_bullet(doc, "Robustesse asymétrique",
        "DashboardService a désormais un timeout (corrigé pendant ce rapport) mais ne capture "
        "toujours pas les erreurs réseau, contrairement à ForecastService — une panne de "
        "data-api produit une erreur 500 non gérée plutôt qu'un message explicite.")
    add_bullet(doc, "Surveillance",
        "Bloc4_app n'est pas instrumenté Prometheus, contrairement à data-api et ml-api — "
        "aucune métrique de trafic ou d'erreur n'est actuellement collectée sur le webapp.")
    add_bullet(doc, "Sécurité",
        "pas de scan automatisé de vulnérabilités des dépendances (pip-audit ou équivalent) ; "
        "DEBUG vaut True par défaut si la variable d'environnement n'est pas positionnée "
        "explicitement — seul docker-compose la force à False aujourd'hui.")
    add_bullet(doc, "Livraison continue",
        "contrairement à ml-api, aucune image webapp n'est construite ni publiée depuis la CI "
        "— le déploiement reste manuel (docker compose up --build).")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 10. CONCLUSION
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("10. Conclusion", level=1)
    doc.add_paragraph(
        "Bloc4_app restitue de façon fonctionnelle les deux services précédemment mis en "
        "service (E1, E3) au travers de trois modules testés (56 tests) et intégrés dans une "
        "chaîne de tests continue. La préparation de ce rapport a permis de trouver et corriger "
        "un vrai bug (absence de timeout sur DashboardService, par symétrie avec le correctif "
        "déjà appliqué à ForecastService en E3) et de documenter honnêtement les écarts entre "
        "l'intention initiale de conduite de projet (board Kanban, rituels d'équipe) et la "
        "réalité d'un développement individuel — plutôt que de les dissimuler."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 11. SOURCES CONSULTÉES
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("11. Sources consultées", level=1)
    sources = [
        "Documentation Django — vues, formulaires, authentification (docs.djangoproject.com)",
        "OWASP Top 10 (2021) — owasp.org/Top10",
        "WCAG 2.1 — w3.org/TR/WCAG21",
        "Documentation uv — docs.astral.sh/uv",
        "Rapports professionnels précédents (E1, E2, E3) du même dépôt",
    ]
    for s in sources:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # ANNEXE
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("Annexe — Repères des sources par thématique", level=1)
    doc.add_paragraph(
        "Table de correspondance rapide entre les thématiques de ce rapport et les fichiers "
        "sources concernés, pour faciliter les échanges avec le jury."
    )
    add_styled_table(doc,
        ["Thématique", "Emplacement"],
        [
            ["Spécifications fonctionnelles", "docs/specifications_fonctionnelles.md"],
            ["Architecture technique", "docs/architecture.md"],
            ["Conduite de projet", "docs/methodologie_agile.md"],
            ["Comptes utilisateurs", "Bloc4_app/accounts/{views,forms,models}.py"],
            ["Tableau de bord et graphiques", "Bloc4_app/dashboard/{views,services}.py, templates/dashboard/{index,charts}.html"],
            ["Classification à la demande", "Bloc4_app/forecast/{views,services,forms}.py, templates/forecast/classify.html"],
            ["Correctif DashboardService (timeout)", "Bloc4_app/dashboard/services.py, tests/test_services.py::TestDashboardService"],
            ["Revue OWASP", "Rapport, §5.5"],
            ["Tests", "Bloc4_app/tests/*.py"],
            ["Intégration continue", ".github/workflows/ci.yml (job test-bloc4)"],
            ["Livraison continue", "Bloc4_app/app.Dockerfile, entrypoint.sh, docker-compose.yml (service webapp)"],
        ],
        col_widths=[6, 11],
    )

    doc.save(str(OUTPUT_PATH))
    print(f"Rapport généré : {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
