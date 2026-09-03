"""Génération du rapport professionnel Bloc 2 — Épreuve E3 (C9, C10, C11, C12, C13)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

OUTPUT_PATH = Path(__file__).parent / "Rapport_Bloc3_Mise_En_Service_Modele_IA.docx"


def set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2E74B5")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D9E2F3")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_bullet(doc, title: str, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"• {title} : ")
    run.bold = True
    p.add_run(text)


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        doc.styles[f"Heading {level}"].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # ── PAGE DE GARDE ──
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Rapport professionnel — Mise en service\n"
        "d'un modèle d'intelligence artificielle"
    )
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Titre professionnel Développeur en Intelligence Artificielle — Bloc 2, Épreuve E3\n"
        "Exposition, intégration, monitorage, tests et livraison continue du modèle de "
        "classification de tendance"
    )
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Projet support : Classification de Tendance Crypto\n"
        "Certification : RNCP37827 — Développeur en Intelligence Artificielle\n"
        "Compétences visées : C9, C10, C11, C12, C13\n"
        "Dépôt : crypto-certification (GitHub)"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_page_break()

    # ── SOMMAIRE ──
    doc.add_heading("Sommaire", level=1)
    for item in [
        "1. Introduction",
        "2. Exposer le modèle via une API REST (C9)",
        "3. Intégrer le modèle dans une application (C10)",
        "4. Monitorer le modèle en production (C11)",
        "5. Tester le modèle automatiquement (C12)",
        "6. Construire une chaîne de livraison continue (C13)",
        "7. Plan de démonstration",
        "8. Bilan et limites assumées",
        "9. Conclusion",
        "10. Sources consultées",
        "Annexe — Repères des sources par thématique",
    ]:
        doc.add_paragraph(item)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Ce rapport porte sur la mise en service du modèle retenu à l'issue du rapport E2 "
        "(Veille, benchmark et paramétrage) : XGBoost, entraîné pour classifier la tendance "
        "du Bitcoin (UP / DOWN / STABLE) à horizon J+1 (journalier) ou H+1 (horaire). Le choix "
        "du modèle et sa comparaison avec un LLM généraliste ne sont pas répétés ici — ce "
        "rapport porte spécifiquement sur ce que le référentiel appelle la « mise en service » : "
        "exposition via une API REST, intégration dans une application existante, monitorage en "
        "production, tests automatisés, et chaîne de livraison continue."
    )
    doc.add_paragraph(
        "Le périmètre technique concerné est le service ml-api (Bloc3_ml, FastAPI, port 8002), "
        "consommé par l'application web Django (Bloc4_app, port 8090). Les compétences évaluées "
        "sont C9 (API REST), C10 (intégration applicative), C11 (monitorage), C12 (tests "
        "automatisés) et C13 (livraison continue)."
    )
    doc.add_paragraph(
        "Comme pour le rapport E2, ce document documente le projet tel qu'il est réellement — "
        "y compris deux anomalies trouvées et corrigées pendant la rédaction (un bug d'intégration "
        "sans garde-fou réseau, et des tests obsolètes ne correspondant plus à l'API réelle) — "
        "plutôt que de ne présenter qu'une version idéalisée du résultat final."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 2. API REST — C9
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("2. Exposer le modèle via une API REST (C9)", level=1)

    doc.add_heading("2.1 Routes exposées", level=2)
    add_styled_table(doc,
        ["Méthode", "Route", "Auth", "Fonction"],
        [
            ["GET", "/", "aucune", "Message d'accueil, pointe vers /docs"],
            ["GET", "/health", "aucune", "Health check (readiness/liveness)"],
            ["GET", "/metrics", "aucune", "Métriques Prometheus (prometheus_fastapi_instrumentator)"],
            ["POST", "/api/v1/authentification/login", "aucune", "Authentification, retourne un JWT (30 min)"],
            ["POST", "/api/v1/classify/classify_daily", "JWT", "Classification à la demande, horizon J+1"],
            ["POST", "/api/v1/classify/classify_hourly", "JWT", "Classification à la demande, horizon H+1"],
        ],
        col_widths=[2, 6.5, 2, 6.5],
    )
    doc.add_paragraph(
        "L'absence d'authentification sur /, /health et /metrics est un choix délibéré, pas un "
        "oubli : ce sont des routes de supervision (utilisées par Docker healthcheck et par le "
        "scrape Prometheus) qui doivent rester accessibles sans jeton pour permettre le monitorage "
        "lui-même — les protéger casserait le health check et le scraping."
    )

    doc.add_heading("2.2 Authentification et documentation", level=2)
    add_bullet(doc, "Mécanisme", "JWT (python-jose, HS256), obtenu via /login avec un couple username/password vérifié contre une valeur unique côté serveur (SecretSettings.E3_PASSWORD) — pas de table utilisateurs à ce stade, un seul compte de service.")
    add_bullet(doc, "Vérification", "get_current_user() (src/api/utils/deps.py) décode et valide le token sur chaque appel aux routes de classification via Depends().")
    add_bullet(doc, "Documentation", "générée nativement par FastAPI au format OpenAPI 3 (/docs Swagger UI, /redoc, /openapi.json) — garantit un standard de marché plutôt qu'une documentation ad hoc, couvre les schémas de requête/réponse et l'authentification par Bearer token.")

    doc.add_heading("2.3 Sécurité — revue OWASP Top 10", level=2)
    doc.add_paragraph(
        "Aucun audit de sécurité formalisé par un tiers n'a été mené. Par souci de rigueur "
        "méthodologique, une auto-évaluation contre les dix catégories du classement OWASP a été "
        "réalisée pour identifier ce qui est déjà couvert par l'architecture actuelle de ml-api et "
        "ce qui reste à renforcer :"
    )
    add_styled_table(doc,
        ["Catégorie OWASP", "Statut sur ml-api"],
        [
            ["A01 — Contrôle d'accès défaillant", "Couvert : JWT requis via Depends(get_current_user) sur les routes de classification ; routes de supervision ouvertes délibérément (cf. §2.1)."],
            ["A02 — Défaillances cryptographiques", "Partiel : JWT signé HS256 avec secret partagé côté serveur, transmis en en-tête HTTP ; le chiffrement en transit (TLS) dépend de l'environnement de déploiement, pas géré en interne à l'API."],
            ["A03 — Injection", "Couvert : aucune requête SQL directe côté ml-api (accès aux données via l'API Bloc1) ; validation des entrées par les schémas Pydantic (ClassifyRequest)."],
            ["A04 — Conception non sécurisée", "Partiel : un seul compte de service partagé, pas de distinction de rôles (contrairement à l'API Bloc1, qui a des rôles lecture seule/admin) — un client authentifié a accès à toutes les routes de classification."],
            ["A05 — Mauvaise configuration de sécurité", "Partiel : pas de rate-limiting, pas de politique CORS explicite documentée — axes d'amélioration identifiés."],
            ["A06 — Composants vulnérables ou obsolètes", "Partiel : dépendances figées via uv.lock, mais pas d'outil de scan de vulnérabilités automatisé intégré à la CI."],
            ["A07 — Identification et authentification défaillantes", "Couvert : mot de passe vérifié côté serveur, JWT à expiration courte (30 min), tests dédiés (401 sur token absent/invalide, cf. §5)."],
            ["A08 — Manquement à l'intégrité des données et logiciels", "Couvert : image Docker construite et taguée par commit dans la CI (cf. §6), pas d'exécution de code non vérifié."],
            ["A09 — Carences des systèmes de contrôle et de journalisation", "Partiel : les requêtes HTTP sont tracées par Prometheus (cf. §4), mais aucune alerte ou journalisation dédiée spécifiquement aux tentatives d'authentification invalides."],
            ["A10 — Falsification de requête côté serveur (SSRF)", "Non applicable : ml-api n'effectue pas de requêtes sortantes pilotées par une entrée utilisateur (les appels vers Bloc1 utilisent des URLs fixes de configuration)."],
        ],
        col_widths=[5, 11],
    )
    doc.add_paragraph(
        "Deux axes de renforcement ressortent clairement de cette revue (limitation de débit, "
        "distinction de rôles) — repris dans le bilan de fin de rapport plutôt que présentés comme "
        "déjà traités."
    )

    doc.add_heading("2.4 Tests de l'API", level=2)
    doc.add_paragraph(
        "13 tests (pytest + TestClient FastAPI) couvrent les 3 routes authentifiées/sensibles "
        "(login, classify_daily, classify_hourly) et les 2 routes de supervision (health, root) :"
    )
    add_bullet(doc, "test_api_auth.py (5 tests)", "login réussi/mot de passe incorrect/mot de passe vide, accès à classify sans token / avec token invalide (401 dans les deux cas).")
    add_bullet(doc, "test_api_classify.py (8 tests)", "health, root, classify_daily/hourly en succès nominal, gestion des champs inconnus ignorés silencieusement par Pydantic, corps de requête manquant (422), et deux cas de paire/modèle introuvable (404 chacun : modèle non entraîné, puis paire inconnue de l'API Bloc1) — anciennement une 500 non interceptée, corrigée suite à un incident réel (cf. rapport E5).")
    doc.add_paragraph(
        "Point méthodologique rencontré en écrivant ces tests : un mock de modèle XGBoost "
        "(unittest.mock.MagicMock) non configuré sur model.get_booster().feature_names se "
        "comporte différemment d'un vrai modèle — MagicMock répond « vrai » à un test de vérité "
        "et s'itère comme une séquence vide, ce qui réduit silencieusement la liste des colonnes "
        "de features à zéro et fait échouer la prédiction sans lever d'exception. Corrigé en "
        "configurant explicitement feature_names=None sur le mock, pour reproduire fidèlement le "
        "comportement d'un modèle non entraîné plutôt que le comportement accidentel d'un mock nu."
    )
    doc.add_paragraph(
        "Suite exécutée : 13 tests réussis (sur les 39 tests au total du dépôt Bloc3_ml, détaillés "
        "au §5). Sources versionnées et accessibles depuis le dépôt Git distant (GitHub)."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 3. INTÉGRATION APPLICATIVE — C10
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("3. Intégrer le modèle dans une application (C10)", level=1)

    doc.add_paragraph(
        "L'application de démonstration est le webapp Django existant (Bloc4_app), qui propose "
        "une page « Classification à la demande » (forecast/classify.html) consommant ml-api via "
        "un service dédié (forecast/services.py::ForecastService)."
    )

    doc.add_heading("3.1 Séquence d'intégration", level=2)
    add_bullet(doc, "Authentification", "ForecastService._get_token() s'authentifie auprès de ml-api avec un compte de service (identifiants en variables d'environnement), cache le token en mémoire pour la durée de la requête.")
    add_bullet(doc, "Appel du modèle", "get_classification() POST vers /classify/classify_daily ou /classify_hourly selon le choix utilisateur, avec le token en en-tête Bearer.")
    add_bullet(doc, "Restitution", "la vue (forecast/views.py::classify_view) affiche la prédiction (label, confiance) ou un message d'erreur explicite via le framework de messages Django, sans jamais laisser une exception non gérée remonter à l'utilisateur.")

    doc.add_heading("3.2 Anomalies trouvées et corrigées pendant ce rapport", level=2)
    doc.add_paragraph(
        "En relisant ForecastService avant rédaction, deux points ont été identifiés et corrigés :"
    )
    add_bullet(doc, "Absence de timeout (bug réel)", "les deux appels requests.post() (login, classification) n'avaient aucun timeout — un ml-api qui répond lentement plutôt que de refuser la connexion pouvait bloquer indéfiniment le thread Django. Corrigé par l'ajout d'un timeout explicite (10s) sur chaque appel.")
    add_bullet(doc, "Renouvellement du token (fragilité, pas un bug vivant)", "le token mis en cache n'était jamais explicitement rafraîchi après expiration (30 min côté ml-api). Vérification faite : ForecastService est instanciée à chaque requête HTTP (forecast/views.py), donc ce cache ne survit jamais d'une requête à l'autre en usage normal — ce n'était pas un bug observable en production actuelle. Renforcé malgré tout par un retry explicite sur 401 (un seul essai, pour ne pas masquer un échec d'authentification durable) plutôt que de dépendre implicitement de ce détail d'instanciation, qui pourrait changer lors d'un futur refactor."
    )
    doc.add_paragraph(
        "Ce choix de documenter une fragilité renforcée par prudence, séparément d'un bug "
        "réellement corrigé, reflète la distinction réelle entre les deux constats plutôt que de "
        "présenter les deux comme des « bugs corrigés » de façon indifférenciée."
    )

    doc.add_heading("3.3 Tests d'intégration", level=2)
    doc.add_paragraph(
        "56 tests couvrent l'application Bloc4_app dans son ensemble (comptes utilisateurs, "
        "dashboard, monitorage, intégration ml-api) :"
    )
    add_styled_table(doc,
        ["Fichier", "Tests", "Périmètre"],
        [
            ["tests/test_accounts.py", "14", "Login, inscription, déconnexion, contrôle d'accès (redirection si non authentifié)"],
            ["tests/test_services.py", "15", "DashboardService (Bloc1) et ForecastService (Bloc3) — dont 2 tests sur le retry ForecastService (401, erreur réseau) et 2 tests ajoutés depuis sur le timeout de DashboardService (cf. rapport E4)"],
            ["tests/test_forecast.py", "6", "Vue de classification à la demande : succès, erreur API, erreur de connexion, formulaire invalide, contrôle d'accès"],
            ["tests/test_dashboard.py", "11", "Dashboard, page de monitorage (cf. §4), graphiques, API interne de données de graphique"],
            ["tests/test_metrics.py", "10", "Calcul de dérive du modèle (fonctions pures, cf. §4) — nouveau module"],
        ],
        col_widths=[6, 2.5, 8.5],
    )
    doc.add_paragraph(
        "Suite exécutée : 56 tests réussis, couverture 97 % (mesurée par pytest-cov, module "
        "dashboard/metrics.py à 100 %). Sources versionnées et accessibles depuis le dépôt Git "
        "distant de l'application."
    )

    doc.add_heading("3.4 Choix d'architecture identifié : lecture en base vs recalcul à la demande", level=2)
    doc.add_paragraph(
        "En instrumentant chaque étape de l'appel « Classifier » pour ce rapport, un déséquilibre "
        "coût/bénéfice est apparu et mérite d'être documenté plutôt que laissé tel quel sans "
        "justification."
    )
    add_styled_table(doc,
        ["Étape (classify_hourly)", "Temps mesuré"],
        [
            ["fetch_recent_ohlcv() — récupère l'historique complet (58 449 lignes)", "2,79 s"],
            ["build_features_for_prediction() — recalcule tous les indicateurs sur cet historique", "0,93 s"],
            ["load_model() — désérialisation du modèle", "0,89 s"],
            ["Total mesuré (hors aller-retours HTTP et double authentification JWT)", "≈ 4,6 s"],
        ],
        col_widths=[13, 4],
    )
    doc.add_paragraph(
        "Or l'API Bloc1 met à jour les OHLCV et le pipeline ML envoie ses prédictions "
        "au maximum une fois par heure (cf. crontabs, §8) : dans les conditions normales, un "
        "recalcul à la demande vise donc la même donnée cible que la dernière prédiction déjà "
        "stockée (predictions_hourly) — vérifié en confrontant les deux valeurs pour un même "
        "horodatage, qui coïncident exactement une fois la bonne paire de trading comparée des "
        "deux côtés. Le bouton « Classifier » recalcule ainsi, à un coût non négligeable, un "
        "résultat que la base connaît déjà la plupart du temps."
    )
    doc.add_paragraph(
        "Le recalcul à la demande garde toutefois un intérêt réel dans deux cas précis : (1) filet "
        "de secours quand le pipeline automatique a pris du retard ou n'a pas tourné — situation "
        "concrètement rencontrée pendant ce rapport (cf. §8, retard du cron journalier) où seule "
        "une prédiction recalculée à la demande reflète une donnée à jour ; (2) démonstration, "
        "notamment pour cette certification, que l'API expose bien une fonction d'inférence "
        "synchrone (C9) et pas seulement un pipeline batch."
    )
    doc.add_paragraph(
        "Recommandation retenue pour une prochaine itération (non implémentée à ce stade, "
        "documentée comme axe d'architecture plutôt que comme correctif urgent) : afficher par "
        "défaut la dernière prédiction stockée (lecture rapide, cohérente avec ce qu'affiche déjà "
        "le dashboard), et réserver le recalcul à un bouton explicite « Recalculer maintenant » "
        "pour les cas où la fraîcheur de la donnée stockée est incertaine. Ce choix réduirait "
        "aussi le risque de confusion entre deux valeurs légèrement différentes pour une même "
        "échéance, observé pendant l'instrumentation de cette section."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 4. MONITORAGE — C11
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("4. Monitorer le modèle en production (C11)", level=1)
    doc.add_paragraph(
        "Le monitorage couvre deux volets distincts, à ne pas confondre l'un avec l'autre : le "
        "monitorage applicatif (la santé de l'API — débit, erreurs, latence) et le monitorage de "
        "la dérive du modèle (la qualité réelle des prédictions dans le temps, comparées à "
        "l'issue effectivement observée)."
    )

    doc.add_heading("4.1 Monitorage applicatif — Prometheus / Grafana", level=2)
    add_bullet(doc, "Collecte", "prometheus_fastapi_instrumentator expose /metrics sur data-api et ml-api ; Prometheus scrape les deux toutes les 15s (monitoring/prometheus.yml).")
    add_bullet(doc, "Alertes", "3 règles configurées (monitoring/alerts.yml) : HighErrorRate (taux 5xx > 5 % sur 2 min), HighLatency (p95 > 5s sur 3 min), ServiceDown (up == 0 sur 1 min).")
    add_bullet(doc, "Restitution", "dashboard Grafana provisionné automatiquement (« API Monitoring - Crypto Certification »), 6 panneaux : requêtes/s par endpoint, taux d'erreur 5xx, latence p50/p95/p99, statut up/down des services, tables de requêtes par endpoint pour data-api et ml-api.")
    doc.add_paragraph(
        "Incident de déploiement rencontré et documenté tel quel : au redémarrage du stack pour "
        "la rédaction de ce rapport, le volume de données Grafana préexistait déjà (créé avant "
        "cette session) avec un mot de passe administrateur différent de celui déclaré en "
        "variable d'environnement — cette dernière ne s'applique qu'à la toute première "
        "initialisation d'une instance Grafana, pas aux démarrages suivants sur un volume déjà "
        "initialisé. Résolu par une réinitialisation officielle et non destructive "
        "(grafana-cli admin reset-admin-password), sans perte de la configuration ni des "
        "dashboards provisionnés."
    )

    doc.add_heading("4.2 Monitorage de la dérive du modèle", level=2)
    doc.add_paragraph(
        "Ce volet répond à l'autre moitié du critère C11 (« métriques courantes ET spécifiques au "
        "projet ») : le monitorage applicatif ci-dessus ne dit rien de la qualité des prédictions "
        "elles-mêmes. Un nouveau module (Bloc4_app/dashboard/metrics.py) et une page dédiée "
        "(/dashboard/monitoring/) ont été construits à cet effet."
    )
    doc.add_paragraph(
        "Choix d'implémentation : plutôt que de créer une nouvelle table d'historique, le module "
        "réutilise deux endpoints déjà exposés par l'API Bloc1 — les prédictions déjà stockées "
        "(predictions_hourly/daily, alimentées par le pipeline Bloc3_ml à chaque cycle "
        "d'entraînement) et les cours OHLCV réels (ohlcv_hourly/daily) — indexés sur les mêmes "
        "clés (trading_pair_id, date). Le label réellement observé est recalculé à partir de la "
        "variation de clôture (même seuil de 0,5 % que l'entraînement, "
        "Bloc3_ml/config/ml_config.yaml) et comparé au label prédit stocké."
    )
    doc.add_paragraph(
        "Résultat obtenu sur des données de production réelles (paire BTC/USDT, granularité "
        "horaire, 21 prédictions évaluables au moment du test) :"
    )
    add_styled_table(doc,
        ["Modèle", "Prédictions évaluées", "Accuracy", "Direction accuracy"],
        [
            ["xgboost", "21", "95,2 %", "0 %"],
        ],
        col_widths=[4, 4, 4, 4],
    )
    doc.add_paragraph(
        "Cet écart spectaculaire n'est pas une anomalie de mesure — direction accuracy exclut "
        "volontairement les cas où prédiction et réalisé valent tous deux STABLE (même définition "
        "que compute_metrics() dans le pipeline d'entraînement, cf. §5). Il confirme en "
        "production, sur des données jamais vues à l'entraînement, exactement le même phénomène "
        "déjà détecté au backtest du rapport E2 : le modèle sur-prédit largement la classe "
        "STABLE, ce qui gonfle artificiellement une métrique d'accuracy brute sans que le modèle "
        "sache réellement détecter une vraie hausse ou baisse. C'est le résultat de monitorage le "
        "plus important de ce rapport — la convergence entre backtest (E2) et production (E3) "
        "renforce la fiabilité du diagnostic plutôt que de l'affaiblir."
    )

    doc.add_heading("4.3 Accessibilité et procédure d'installation", level=2)
    add_bullet(doc, "Accessibilité", "page de monitorage construite avec les composants Bootstrap natifs déjà utilisés dans le reste de l'application (tableaux, boutons), héritant de la navigation clavier et du contraste du thème existant.")
    add_bullet(doc, "Installation", "aucune dépendance supplémentaire — dashboard/metrics.py est un module Python pur (aucun appel réseau propre), les données transitent par DashboardService déjà présent.")
    add_bullet(doc, "Test avant mise en service", "le module de calcul a été testé isolément (10 tests unitaires purs, cf. §5) avant d'être branché à la vue Django, conformément à une logique de bac à sable avant intégration.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 5. TESTS AUTOMATISÉS — C12
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("5. Tester le modèle automatiquement (C12)", level=1)
    doc.add_paragraph(
        "95 tests automatisés au total couvrent le modèle et son intégration : 39 dans "
        "Bloc3_ml (pipeline de features, métriques, API) et 56 dans Bloc4_app (application, "
        "dont le nouveau monitorage de dérive)."
    )

    doc.add_heading("5.1 Tests du pipeline de features et du modèle (Bloc3_ml)", level=2)
    add_styled_table(doc,
        ["Fichier", "Tests", "Périmètre"],
        [
            ["tests/test_api_auth.py", "5", "Authentification JWT (cf. §2.4)"],
            ["tests/test_api_classify.py", "8", "Endpoints de classification (cf. §2.4)"],
            ["tests/test_build_features.py", "19", "Construction des features : lags, rendements, indicateurs techniques, features temporelles, cible ; dont 2 tests dédiés à l'absence de fuite de données (§5.2)"],
            ["tests/test_evaluate_model.py", "7", "Fonction compute_metrics() : predictions parfaites/toutes fausses, présence des clés, plage de valeurs, direction accuracy avec/sans cas STABLE-STABLE, liste vide"],
        ],
        col_widths=[6, 2.5, 8.5],
    )

    doc.add_heading("5.2 Absence de fuite de données (data leakage)", level=2)
    doc.add_paragraph(
        "Plutôt que de supposer que les décalages temporels (lags, cible) sont corrects, chaque "
        "colonne est recalculée manuellement à partir des données brutes et comparée à la colonne "
        "produite par le pipeline — un décalage cassé ou inversé ferait échouer ces tests :"
    )
    add_bullet(doc, "test_all_lag_columns_match_manual_shift", "toutes les colonnes de lag (close, volume, high, low × 5 décalages) recalculées via pandas.Series.shift() indépendamment du code testé, comparées ligne à ligne avec pandas.testing.assert_series_equal.")
    add_bullet(doc, "test_target_reflects_next_period_not_current_or_past", "la cible est recalculée manuellement à partir de la variation de clôture entre t et t+1 (pas t-1 et t), comparée ligne à ligne à la sortie réelle de add_target() — vérifie explicitement que le modèle apprend à prédire l'avenir et non à mémoriser le passé. Vérifie également que la toute dernière ligne du jeu de données a une cible NaN par construction (aucune période suivante connue), un comportement volontaire du pipeline (nécessaire pour que l'inférence dispose d'une ligne de features sans devoir connaître le futur) plutôt qu'un défaut.")

    doc.add_heading("5.3 Couverture de code (Bloc3_ml)", level=2)
    doc.add_paragraph(
        "Mesurée avec pytest-cov, sur l'ensemble du dépôt Bloc3_ml (454 lignes instrumentées, "
        "54 % de couverture globale) :"
    )
    add_styled_table(doc,
        ["Module", "Couverture", "Commentaire"],
        [
            ["src/features/build_features.py", "100 %", "Pipeline de features, entièrement couvert (dont anti-leakage)"],
            ["src/api/utils/auth.py, login.py, deps.py, classes.py", "100 %", "Authentification et schémas de requête"],
            ["src/api/routes/classify.py", "94 %", "Endpoints de classification"],
            ["src/api/api.py", "91 %", "App FastAPI (bloc __main__ non couvert, normal)"],
            ["src/model/evaluate_model.py", "51 %", "compute_metrics() couvert ; evaluate_classifier() (boucle d'entraînement/évaluation complète) non couvert par des tests unitaires — nécessiterait un jeu de données synthétique multi-périodes, hors périmètre de cette itération"],
            ["src/model/train_model.py", "30 %", "Partiellement couvert indirectement via les tests d'API (mocks)"],
            ["src/api/utils/functions.py", "63 %", "load_model/fetch_recent_ohlcv couverts via les mocks des tests API, y compris les deux chemins 404 ajoutés (cf. §2.4) ; build_features_for_prediction non exercé directement"],
            ["src/data/*, src/model/initiate_classifier.py, predict_model.py, save_model.py, src/monitoring/monitor_training.py, src/utils/classes.py", "0 %", "Code d'orchestration du pipeline d'entraînement (appels réseau vers Bloc1, écriture de modèles sur disque) — non couvert, choix de périmètre assumé pour cette itération (cf. bilan)"],
        ],
        col_widths=[6.5, 2.5, 8],
    )
    doc.add_paragraph(
        "Ce choix de périmètre n'est pas un oubli mais une priorisation assumée : le code testé "
        "en premier est celui qui présente le risque le plus direct pour la qualité des "
        "prédictions (features, métriques, API), avant le code d'orchestration qui appelle des "
        "services externes et nécessiterait des mocks plus lourds pour un gain de couverture "
        "moindre en termes de risque réel."
    )

    doc.add_heading("5.4 Installation et exécution", level=2)
    doc.add_paragraph("cd Bloc3_ml && uv sync --dev", style="Intense Quote")
    doc.add_paragraph(
        "pytest tests/ -v --cov=src --cov-report=term-missing",
        style="Intense Quote",
    )
    doc.add_paragraph(
        "cd Bloc4_app && uv sync --dev && pytest tests/ -v --cov=. --cov-report=term-missing",
        style="Intense Quote",
    )
    doc.add_paragraph(
        "Ces commandes sont celles effectivement exécutées par la CI (cf. §6). Sources "
        "versionnées et accessibles depuis le dépôt Git distant (GitHub)."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 6. LIVRAISON CONTINUE — C13
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("6. Construire une chaîne de livraison continue (C13)", level=1)
    doc.add_paragraph(
        "La chaîne (GitHub Actions, .github/workflows/ci.yml) couvrait déjà les tests et le lint "
        "avant ce rapport (jobs test-bloc1, test-bloc3, test-bloc4, lint, déclenchés sur push/pull "
        "request vers main/develop). Deux étapes manquaient pour répondre pleinement au critère "
        "C13 (« packaging et déploiement ») : elles ont été ajoutées et testées localement."
    )

    doc.add_heading("6.1 Étapes de la chaîne", level=2)
    add_styled_table(doc,
        ["Job", "Déclencheur", "Rôle"],
        [
            ["test-bloc3", "push / pull_request sur main, develop", "Exécute les 39 tests Bloc3_ml avec couverture (préexistant)"],
            ["build-ml-api", "push sur main, après succès de test-bloc3", "Construit l'image Docker ml-api (api.Dockerfile) et la publie sur ghcr.io, taguée latest et par SHA de commit"],
            ["deploy-ml-api", "push sur main, après succès de build-ml-api", "Vérifie que l'image publiée est réellement récupérable (docker pull), documente la commande de redéploiement"],
        ],
        col_widths=[4, 6, 7],
    )
    doc.add_paragraph(
        "La dépendance entre étapes (needs: dans la syntaxe GitHub Actions, équivalent du "
        "depends_on d'autres moteurs de CI) garantit qu'aucune image n'est publiée si les tests "
        "échouent, et qu'aucun déploiement n'est tenté si la construction échoue."
    )

    doc.add_heading("6.2 Packaging", level=2)
    doc.add_paragraph(
        "L'image est construite via docker/build-push-action, authentifiée sur le registre par le "
        "GITHUB_TOKEN natif (aucun secret supplémentaire à gérer), et poussée vers "
        "ghcr.io/<owner>/crypto-ml-api avec deux tags : latest (dernière version sur main) et le "
        "SHA du commit (traçabilité exacte de la version déployée). Le build a été testé "
        "localement avec exactement le même contexte que la CI (Bloc3_ml/api.Dockerfile) avant "
        "d'être intégré au workflow, pour éviter de découvrir une erreur de build seulement après "
        "un push."
    )

    doc.add_heading("6.3 Déploiement — périmètre assumé", level=2)
    doc.add_paragraph(
        "Ce projet n'a pas de serveur de production dédié (contrairement à un contexte "
        "d'entreprise avec infrastructure déployée) — le déploiement est donc scopé à ce qui est "
        "réellement vérifiable plutôt que simulé. Le job deploy-ml-api vérifie que l'image "
        "fraîchement publiée est effectivement récupérable depuis le registre (docker pull), ce "
        "qui constitue une preuve concrète que le packaging produit un artefact utilisable — puis "
        "documente la commande réelle de redéploiement plutôt que de prétendre exécuter un "
        "déploiement automatisé vers un serveur qui n'existe pas :"
    )
    doc.add_paragraph(
        "docker compose pull ml-api && docker compose up -d ml-api",
        style="Intense Quote",
    )
    doc.add_paragraph(
        "Ce choix est documenté explicitement comme un périmètre assumé (cf. bilan) plutôt que "
        "présenté comme une limitation cachée."
    )

    doc.add_heading("6.4 Installation et vérification de la chaîne", level=2)
    add_bullet(doc, "Secrets", "aucun secret à créer manuellement — le GITHUB_TOKEN fourni automatiquement par GitHub Actions suffit pour l'authentification au registre ghcr.io (permissions packages: write déclarées dans le job).")
    add_bullet(doc, "Vérification locale avant intégration", "docker build -f Bloc3_ml/api.Dockerfile -t crypto-ml-api:ci-test Bloc3_ml — build réussi en ~19s, confirmant que le Dockerfile fonctionne avant de dépendre du runner GitHub.")
    add_bullet(doc, "Syntaxe", "fichier YAML validé (python3 -c \"import yaml; yaml.safe_load(...)\") avant intégration.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 7. PLAN DE DÉMONSTRATION
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("7. Plan de démonstration", level=1)
    doc.add_paragraph(
        "Ce que je prévois de montrer en direct lors de la soutenance, pour chacun des "
        "composants attendus par l'épreuve."
    )

    doc.add_heading("1. L'API du modèle", level=2)
    add_bullet(doc, "Documentation interactive", "ouverture de /docs, présentation des routes de classification et d'authentification.")
    add_bullet(doc, "Authentification", "login via /api/v1/authentification/login, puis appel de classify_daily avec le token obtenu.")
    add_bullet(doc, "Erreurs", "appel sans token puis avec un token invalide — démonstration du 401 dans les deux cas.")
    add_bullet(doc, "Tests", "exécution en direct de pytest tests/ -v sur Bloc3_ml.")

    doc.add_heading("2. L'application enrichie par l'API", level=2)
    add_bullet(doc, "Classification à la demande", "sélection d'une paire, affichage de la prédiction obtenue via ml-api.")
    add_bullet(doc, "Monitorage de dérive", "page /dashboard/monitoring/, lecture des métriques accuracy/direction accuracy en conditions réelles.")
    add_bullet(doc, "Tests", "exécution en direct de pytest tests/ -v sur Bloc4_app.")

    doc.add_heading("3. Le monitorage applicatif et la livraison continue", level=2)
    add_bullet(doc, "Grafana", "dashboard « API Monitoring », déclenchement d'un appel de classification en direct pour montrer le panneau requêtes/s bouger.")
    add_bullet(doc, "CI/CD", "présentation du workflow GitHub Actions, de l'enchaînement test → build → deploy, et de l'image publiée sur ghcr.io.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 8. BILAN ET LIMITES ASSUMÉES
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("8. Bilan et limites assumées", level=1)
    doc.add_paragraph(
        "Par souci d'honnêteté professionnelle, les limites déjà évoquées au fil du rapport sont "
        "rassemblées ici plutôt que laissées dispersées :"
    )
    add_bullet(doc, "Pas de distinction de rôles sur ml-api", "un seul compte de service partagé donne accès à toutes les routes de classification (contrairement à l'API Bloc1, qui distingue lecture seule et admin) — identifié dans la revue OWASP (§2.3), à traiter dans une prochaine itération si l'API devait être exposée à plusieurs clients aux besoins différents.")
    add_bullet(doc, "Pas de rate-limiting ni de politique CORS documentée", "sur ml-api — axe de renforcement identifié en §2.3, non bloquant pour un usage interne actuel.")
    add_bullet(doc, "Couverture de tests partielle sur le code d'orchestration", "le pipeline d'entraînement (appels réseau, écriture de modèles) n'est pas couvert par des tests unitaires (cf. §5.3) — un choix de périmètre assumé, pas un oubli.")
    add_bullet(doc, "Pipeline de données dépendant de la disponibilité de la machine hôte", "les mises à jour OHLCV et les prédictions du pipeline (Bloc1/Bloc3) sont planifiées par cron à l'intérieur des conteneurs Docker. Sur un environnement de développement qui n'est pas actif 24/7 (mise en veille de la machine), les créneaux manqués ne sont pas rattrapés automatiquement — constaté concrètement pendant la rédaction de ce rapport (données journalières bloquées plusieurs semaines faute de fenêtre de déclenchement atteinte, écart de plusieurs heures côté horaire après une nuit de veille). C'est une limite inhérente à l'hébergement local, qui disparaîtrait sur un serveur de production toujours actif — corrigé ponctuellement par un déclenchement manuel du script de mise à jour, documenté ici plutôt que masqué.")
    add_bullet(doc, "Déploiement scopé au registre d'images", "pas de serveur de production réel à ce stade (cf. §6.3) — le déploiement automatisé s'arrête à la vérification que l'image est publiée et récupérable, la mise en service effective sur un environnement cible restant une commande manuelle documentée.")
    add_bullet(doc, "Qualité prédictive du modèle", "confirmée en production comme au backtest (§4.2) : le modèle sur-prédit la classe STABLE, avec une direction accuracy proche de 0 % sur l'échantillon observé — un axe d'amélioration prioritaire pour une prochaine itération du modèle (rééquilibrage des classes ou ajustement du seuil de classification), documenté plutôt que dissimulé derrière l'accuracy brute qui, seule, donnerait une impression trompeuse de performance.")
    add_bullet(doc, "Classification à la demande recalculée plutôt que lue en base", "le bouton « Classifier » recalcule systématiquement une prédiction (≈4,6 s, cf. §3.4) alors que le résultat coïncide la plupart du temps avec la dernière valeur déjà stockée — la donnée sous-jacente n'étant rafraîchie qu'une fois par heure par le pipeline. Une lecture en base par défaut, avec recalcul en option explicite, serait plus économe sans perte de fraîcheur dans le cas général ; conservé tel quel pour cette version car le recalcul reste un filet de secours utile en cas de retard du pipeline (cf. point précédent sur le cron) et démontre la fonction d'inférence synchrone de l'API (C9).")
    doc.add_paragraph(
        "Aucun de ces points ne remet en cause le fonctionnement du service tel que mis en "
        "service ; ils constituent les axes d'amélioration prioritaires pour une prochaine "
        "itération."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 9. CONCLUSION
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("9. Conclusion", level=1)
    doc.add_paragraph(
        "Ce rapport documente la mise en service complète du modèle retenu en E2 : une API REST "
        "authentifiée et documentée (C9), intégrée dans une application Django existante avec ses "
        "propres garde-fous réseau (C10), surveillée à deux niveaux — santé applicative et dérive "
        "du modèle lui-même (C11) —, sécurisée par 95 tests automatisés couvrant explicitement "
        "l'absence de fuite de données (C12), et livrée par une chaîne d'intégration et de "
        "livraison continue allant du test au packaging (C13)."
    )
    doc.add_paragraph(
        "Comme pour le rapport E2, le chemin parcouru n'a pas été linéaire : deux anomalies "
        "réelles (timeout manquant, tests obsolètes ne correspondant plus à l'API) ont été "
        "trouvées et corrigées en cours de rédaction, et un résultat inattendu — la faiblesse de "
        "la direction accuracy du modèle, confirmée à la fois par le backtest et par le "
        "monitorage de production — a été documenté sans être minimisé. C'est cette convergence "
        "entre ce qui est écrit dans le rapport et ce qui est réellement vérifiable dans le dépôt "
        "qui, je l'espère, en constitue la valeur la plus solide."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 10. SOURCES
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("10. Sources consultées", level=1)
    sources = [
        "FastAPI — Documentation officielle (OpenAPI, sécurité, dépendances) — docs primaire, éditeur",
        "python-jose — Documentation JWT (source primaire, bibliothèque utilisée)",
        "OWASP — Top 10:2021 (source primaire, référentiel de sécurité)",
        "Prometheus — Documentation officielle (scraping, PromQL, alerting rules)",
        "prometheus-fastapi-instrumentator — Documentation du package (GitHub, source primaire)",
        "Grafana — Documentation officielle (provisioning, grafana-cli, gestion des identifiants admin)",
        "GitHub Actions — Documentation officielle (workflows, GITHUB_TOKEN, docker/build-push-action)",
        "GitHub Container Registry (ghcr.io) — Documentation officielle (authentification, tags)",
        "pytest / pytest-cov / unittest.mock — Documentation officielle",
        "Règlement spécifique RNCP37827 — Développeur en Intelligence Artificielle (Simplon, 2023)",
        "Référentiel Activités Compétences et évaluation — Bloc 2, épreuve E3 (Simplon, 2023)",
        "Sources internes : Bloc3_ml/, Bloc4_app/, monitoring/, .github/workflows/ci.yml, "
        "docs/monitoring_applicatif.md, rapport E2 (Rapport_Bloc2_Veille_Benchmark_Parametrage_IA)",
    ]
    for s in sources:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # ANNEXE
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("Annexe — Repères des sources par thématique", level=1)
    doc.add_paragraph(
        "Table de correspondance rapide entre les thématiques de ce rapport et les fichiers "
        "sources concernés, pour faciliter les échanges avec le jury."
    )
    add_styled_table(doc,
        ["Thématique", "Emplacement"],
        [
            ["Routes de l'API ML", "Bloc3_ml/src/api/routes/{classify,login}.py"],
            ["Authentification JWT", "Bloc3_ml/src/api/utils/{auth,deps}.py"],
            ["Revue OWASP", "Rapport, §2.3"],
            ["Tests de l'API ML", "Bloc3_ml/tests/{test_api_auth,test_api_classify}.py"],
            ["Intégration Django (C10)", "Bloc4_app/forecast/{services,views}.py"],
            ["Tests d'intégration", "Bloc4_app/tests/{test_services,test_forecast}.py"],
            ["Monitorage applicatif", "monitoring/{prometheus.yml,alerts.yml}, monitoring/grafana/"],
            ["Monitorage de dérive du modèle", "Bloc4_app/dashboard/{metrics,views}.py, templates/dashboard/monitoring.html"],
            ["Tests du monitorage de dérive", "Bloc4_app/tests/test_metrics.py, tests/test_dashboard.py::TestMonitoringView"],
            ["Construction des features + anti-leakage", "Bloc3_ml/src/features/build_features.py, tests/test_build_features.py::TestNoDataLeakage"],
            ["Métriques du modèle", "Bloc3_ml/src/model/evaluate_model.py, tests/test_evaluate_model.py"],
            ["Chaîne de livraison continue", ".github/workflows/ci.yml (jobs build-ml-api, deploy-ml-api)"],
            ["Packaging Docker", "Bloc3_ml/api.Dockerfile"],
        ],
        col_widths=[6, 11],
    )

    doc.save(str(OUTPUT_PATH))
    print(f"Rapport généré : {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
