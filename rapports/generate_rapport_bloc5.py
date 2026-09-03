"""Génération du rapport professionnel Bloc 3 — Épreuve E5 (C20, C21).

Format volontairement court (2-5 pages, consigne de l'épreuve) : pas de page de garde
séparée, pas de sommaire, pas de plan de démo ni d'annexe. Structure calquée sur le
rapport de référence (Poutot E5) : introduction, dispositif de monitorage, incident,
diagnostic, résolution, conclusion.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

OUTPUT_PATH = Path(__file__).parent / "Rapport_Bloc5_Monitoring_Incident.docx"
ASSETS_DIR = Path(__file__).parent / "assets_bloc5"


def set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2E74B5")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D9E2F3")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_bullet(doc, title: str, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{title} : ")
    run.bold = True
    p.add_run(text)


def add_figure(doc, filename: str, caption: str, width_cm=11):
    doc.add_picture(str(ASSETS_DIR / filename), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)


def add_code(doc, code: str):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {qn("w:fill"): "F2F2F2", qn("w:val"): "clear"})
    pPr.append(shd)


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(5)
    style.paragraph_format.line_spacing = 1.1

    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(4)

    # ── EN-TÊTE (pas de page de garde séparée) ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Monitorage applicatif et résolution d'un incident technique")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Titre professionnel Développeur en Intelligence Artificielle — Bloc 3, Épreuve E5\n"
        "Projet support : Classification de Tendance Crypto — Compétences visées : C20, C21\n"
        "Dépôt : crypto-certification (GitHub)"
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════
    # INTRODUCTION
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "Ce document présente le dispositif de monitorage applicatif mis en place sur les "
        "APIs data-api (Bloc1) et ml-api (Bloc3), ainsi qu'un incident technique réel rencontré "
        "et résolu au fil de ce travail : un appel à la classification à la demande avec un "
        "symbole de paire sans modèle entraîné (par exemple un nom mal orthographié) faisait "
        "remonter une erreur 500 non gérée, au lieu d'une erreur 404 claire. Les deux sujets "
        "sont traités ensemble : c'est en testant l'API réellement démarrée avec des entrées "
        "invalides, plutôt qu'en se fiant à la seule suite de tests automatisés, que l'anomalie "
        "a été repérée puis confirmée par le monitorage déjà branché."
    )

    # ══════════════════════════════════════════════════════════════════
    # DISPOSITIF DE MONITORING
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("Dispositif de monitorage applicatif", level=1)
    doc.add_paragraph(
        "Prometheus et Grafana, déployés localement via docker-compose, comme le reste du "
        "projet (il n'existe pas d'environnement de production séparé sur ce projet). "
        "Prometheus va lui-même chercher les métriques sur data-api et ml-api toutes les "
        "15 secondes (modèle « pull » : pas de programme supplémentaire à installer sur les "
        "APIs, qui exposent déjà leurs métriques via prometheus-fastapi-instrumentator). "
        "Grafana restitue un dashboard et les règles d'alerte au-dessus de Prometheus, sans "
        "développement spécifique."
    )
    add_bullet(doc, "Métriques", "http_requests_total (nombre de requêtes, ventilé par route, "
        "méthode HTTP et code de réponse) et http_request_duration_seconds (temps de "
        "réponse), exposées sur /metrics par les deux APIs.")
    add_bullet(doc, "Dashboard", "« API Monitoring - Crypto Certification » (6 panels : "
        "requêtes/s, taux d'erreur 5xx, latence p50/p95/p99, statut UP/DOWN, requêtes par "
        "endpoint pour chaque API), provisionné automatiquement depuis monitoring/.")
    add_bullet(doc, "3 règles d'alerte", "(monitoring/alerts.yml) HighErrorRate (plus de 5 % "
        "de requêtes en 5xx sur 5 min), HighLatency (p95 supérieur à 5 s), ServiceDown "
        "(service injoignable). Ces règles sont évaluées par Prometheus, avec un état visible "
        "sur sa page /alerts, mais aucune n'envoie de notification externe (e-mail, Slack...) : "
        "il faut consulter cette page pour savoir si une alerte s'est déclenchée.")
    add_bullet(doc, "Installation", "docker compose up -d prometheus grafana (ports 9090 "
        "et 3000, admin/admin).")
    add_figure(doc, "fig1_prometheus_targets.png",
        "Figure 1 — Cibles Prometheus réelles : data-api et ml-api, toutes UP.")

    # ══════════════════════════════════════════════════════════════════
    # L'INCIDENT
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("L'incident : une paire inconnue fait échouer la classification", level=1)
    doc.add_paragraph(
        "En testant ml-api, réellement démarré, avec des entrées limites (des appels curl "
        "manuels, pas la suite de tests automatisés), l'appel POST /classify/classify_daily "
        "avec {\"trading_pair_symbol\": \"FAKE-COIN\"} (un symbole de paire imaginaire) a "
        "échoué en 500 Internal Server Error, alors qu'un appel équivalent sur une paire "
        "réelle (BTC-USDT) répond normalement. Les logs du conteneur donnent la cause "
        "immédiate :"
    )
    add_code(doc, "FileNotFoundError: Modèle FAKE-COIN introuvable : /app/models/day_models/FAKE-COIN.pkl")
    doc.add_paragraph(
        "Diagnostic (src/api/utils/functions.py) : load_model() et fetch_recent_ohlcv() "
        "lèvent des exceptions Python brutes (FileNotFoundError, Exception) au lieu d'une "
        "HTTPException. FastAPI retombe alors sur son gestionnaire générique, qui répond 500 "
        "pour toute exception non gérée. Un test existant du dépôt attendait même "
        "explicitement le code 500 comme réponse correcte : la suite de tests automatisés ne "
        "pouvait donc pas révéler ce défaut, puisqu'elle avait normalisé le mauvais "
        "comportement plutôt que de vérifier le bon."
    )
    doc.add_paragraph(
        "Détection via le monitorage : dans les 15 secondes qui suivent (le prochain passage "
        "de Prometheus), la requête est enregistrée en 5xx, visible dans le panel « Taux "
        "d'erreur 5xx » et dans la table « Requêtes par endpoint (ml-api) ». Confirmé aussi "
        "par une requête directe à l'API Prometheus, pas seulement sur le rendu Grafana."
    )
    add_figure(doc, "fig3_grafana_incident.png",
        "Figure 2 — Dashboard pendant la reproduction de l'incident : le taux d'erreur 5xx "
        "(bloc3-ml-api) et la dernière ligne de la table de droite (status 5xx).")

    doc.add_heading("Résolution", level=2)
    doc.add_paragraph(
        "Tracé sur l'outil de suivi du dépôt (Issue GitHub #10). Les deux points remontent "
        "désormais une HTTPException(404) explicite avec un message "
        "clair (« Aucun modèle entraîné pour '{symbol}' », « Paire de trading introuvable »). "
        "Le test qui attendait le code 500 a été corrigé pour attendre 404, et un second test a été "
        "ajouté pour couvrir le second point d'entrée du même défaut (paire inconnue de "
        "l'API Bloc1). Suite complète du dépôt Bloc3_ml : 39 tests réussis. Vérifié en direct "
        "après reconstruction du conteneur : FAKE-COIN → 404 avec message clair, BTC-USDT → "
        "200 (comportement nominal inchangé)."
    )
    add_figure(doc, "fig4_grafana_recovery.png",
        "Figure 3 — Retour à la normale : la ligne d'erreur 5xx s'arrête au moment du "
        "correctif et ne reprend pas.")
    doc.add_paragraph(
        "Limite honnête : l'alerte HighErrorRate (ratio sur fenêtre de 5 min, seuil tenu "
        "2 min) est restée « inactive » pendant tout ce test : le correctif a été appliqué en "
        "moins de deux minutes, et le faible volume de requêtes ne suffisait pas à stabiliser "
        "un taux mesurable. Le dashboard, consulté activement, a détecté l'anomalie bien plus "
        "vite que l'alerte n'aurait pu se déclencher."
    )

    # ══════════════════════════════════════════════════════════════════
    # CONCLUSION
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "Le monitorage mis en place a rempli son rôle au-delà de la démonstration : une "
        "requête invalide envoyée pendant la préparation de ce rapport a immédiatement été "
        "visible dans Prometheus et Grafana, et le retour à la normale après correctif l'a été "
        "tout autant. Le correctif est appliqué, testé et vérifié en direct."
    )

    doc.save(str(OUTPUT_PATH))
    print(f"Rapport généré : {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
