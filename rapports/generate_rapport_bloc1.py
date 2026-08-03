"""Génération du rapport professionnel Bloc 1 au format Word (.docx)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

OUTPUT_PATH = Path(__file__).parent / "Rapport_Bloc1_Collecte_Stockage_Donnees.docx"


def set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2E74B5")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.style = doc.styles["Normal"]
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D9E2F3")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # ── PAGE DE GARDE ──
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Rapport Professionnel — Bloc 1")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Réaliser la collecte, le stockage et la mise à disposition\n"
        "des données d'un projet en intelligence artificielle"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Projet : Classification de Tendance Crypto\n"
        "Certification : Développeur en Intelligence Artificielle (RNCP37827)\n"
        "Évaluation : E1 — Mise en situation professionnelle (C1 à C5)"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_page_break()

    # ── SOMMAIRE ──
    doc.add_heading("Sommaire", level=1)
    sommaire_items = [
        "1. Contexte et présentation du projet",
        "2. C1 — Extraction des données depuis des sources multiples",
        "3. C2 / C3 — Requêtes SQL et agrégation des données",
        "4. C4 — Base de données et conformité RGPD",
        "5. C5 — API REST et authentification",
        "Annexes",
    ]
    for item in sommaire_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # PAGE 1 — CONTEXTE
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("1. Contexte et présentation du projet", level=1)

    doc.add_heading("1.1 Objectif du projet", level=2)
    doc.add_paragraph(
        "Le projet « Classification de Tendance Crypto » a pour objectif de développer un système "
        "complet de collecte, de stockage et de mise à disposition de données de marché de crypto-monnaies "
        "(BTC, ETH, etc.), afin d'alimenter un modèle de machine learning capable de classifier "
        "les tendances de prix (hausse, baisse, neutre)."
    )
    doc.add_paragraph(
        "Ce rapport couvre le Bloc 1 de la certification RNCP37827 : la couche données du projet, "
        "depuis l'extraction multi-sources jusqu'à l'exposition via une API REST sécurisée."
    )

    doc.add_heading("1.2 Acteurs et organisation du travail", level=2)
    doc.add_paragraph(
        "Le projet est réalisé en autonomie dans le cadre d'une certification professionnelle, "
        "en suivant une méthodologie agile (sprints itératifs). Le code est versionné sur Git, "
        "organisé en modules correspondant aux compétences du bloc (C1 à C5)."
    )

    doc.add_heading("1.3 Contraintes techniques", level=2)
    add_styled_table(doc,
        ["Contrainte", "Détail"],
        [
            ["Volume de données", "Données OHLCV à la minute, horaire et journalier pour plusieurs paires de trading"],
            ["Multi-sources", "4 sources distinctes : 2 API REST, 1 fichier CSV, 1 scraping web"],
            ["Pérennisation", "Collecte automatisée par cron jobs (horaire et journalier)"],
            ["Sécurité", "Accès aux données protégé par authentification JWT"],
        ],
        col_widths=[4, 12],
    )

    doc.add_heading("1.4 Pile technique", level=2)
    add_styled_table(doc,
        ["Composant", "Technologie"],
        [
            ["Langage", "Python 3.11+"],
            ["API REST", "FastAPI (OpenAPI / Swagger)"],
            ["Base de données", "PostgreSQL 16"],
            ["ORM", "SQLAlchemy"],
            ["Scraping", "Scrapy"],
            ["Authentification", "JWT (python-jose) + bcrypt (passlib)"],
            ["Conteneurisation", "Docker / Docker Compose"],
            ["Tests", "pytest (52 tests)"],
            ["Métriques", "Prometheus (prometheus-fastapi-instrumentator)"],
        ],
        col_widths=[5, 11],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # PAGE 2 — C1 : EXTRACTION
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("2. C1 — Extraction des données depuis des sources multiples", level=1)

    doc.add_paragraph(
        "Conformément aux exigences du référentiel, l'extraction des données est réalisée depuis "
        "un mix de sources couvrant les différents types attendus : service web (API REST), "
        "fichier de données (CSV), scraping de page web, et base de données."
    )

    doc.add_heading("2.1 Sources de données", level=2)
    add_styled_table(doc,
        ["Source", "Type", "Données collectées", "Module"],
        [
            ["Binance", "API REST", "OHLCV minute / horaire / journalier", "C1_extraction/extract_binance_data.py"],
            ["CoinMarketCap", "API REST", "Currencies, paires de trading, classements", "C1_extraction/extract_coinmarketcap.py"],
            ["CryptoDownload", "Fichiers CSV", "Données historiques OHLCV", "C1_extraction/extract_csv_data.py"],
            ["CoinTelegraph", "Scraping (Scrapy)", "Actualités Bitcoin (titre, date, URL, catégorie)", "scraping/spiders/cointelegraph_spider.py"],
        ],
        col_widths=[3, 3, 5.5, 5.5],
    )

    doc.add_heading("2.2 Scraping web — CoinTelegraph", level=2)
    doc.add_paragraph(
        "Le spider Scrapy cible la page des articles Bitcoin de CoinTelegraph. Il respecte les bonnes pratiques : "
        "respect du fichier robots.txt, délai de 2 secondes entre les requêtes, identification via un User-Agent "
        "personnalisé. Un pipeline de déduplication par URL évite les doublons. Les données sont exportées "
        "au format JSON."
    )

    doc.add_heading("2.3 Pérennisation de la collecte", level=2)
    doc.add_paragraph(
        "La collecte est automatisée via des tâches cron exécutées dans un conteneur Docker dédié :"
    )
    add_styled_table(doc,
        ["Fréquence", "Commande", "Description"],
        [
            ["Toutes les heures (XX:02)", "python -m update_ohlcv --frequency hour", "Mise à jour des données horaires depuis Binance"],
            ["Quotidien (00:01)", "python -m update_ohlcv --frequency day", "Mise à jour des données journalières depuis Binance"],
        ],
        col_widths=[4, 6.5, 5.5],
    )

    doc.add_heading("2.4 Gestion des erreurs", level=2)
    doc.add_paragraph(
        "Chaque script intègre : l'initialisation des dépendances et des connexions externes, "
        "les règles logiques de traitement, la gestion des exceptions (try/except avec logging), "
        "et la sauvegarde des résultats. Les insertions en batch incluent un fallback individuel "
        "en cas de conflit d'unicité, et les lignes échouées sont journalisées dans des fichiers JSON "
        "dédiés (data/logs/)."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # PAGE 3 — C2 / C3 : SQL + AGRÉGATION
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("3. C2 / C3 — Requêtes SQL et agrégation des données", level=1)

    doc.add_heading("3.1 C2 — Requêtes SQL (SQLAlchemy ORM)", level=2)
    doc.add_paragraph(
        "Les requêtes d'extraction de données sont développées via SQLAlchemy ORM dans le module C2_query/. "
        "Chaque entité dispose de fonctions de requête dédiées, documentées et optimisées."
    )

    doc.add_paragraph("Optimisations mises en œuvre :")
    optimizations = [
        ("Jointures eager (joinedload)", "Chargement anticipé des relations TradingPair → Currency pour éviter le problème N+1 queries."),
        ("Filtrage par date", "Validation côté API (format YYYY-MM-DD) avant exécution de la requête SQL."),
        ("Tri ascendant", "Tri par date pour garantir l'affichage chronologique des données."),
        ("Contraintes d'unicité", "Clés composites (trading_pair_id, date) pour empêcher les doublons en base."),
        ("Insertion batch", "bulk_insert_mappings par lots de 10 000 lignes avec fallback individuel en cas de conflit."),
    ]
    for title_text, desc in optimizations:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title_text} : ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)

    doc.add_heading("3.2 C3 — Agrégation et nettoyage des données", level=2)
    doc.add_paragraph(
        "Les données brutes proviennent de sources aux formats hétérogènes. Le script d'agrégation "
        "(C3_aggregate_ohlcv/aggregate_ohlcv.py) les consolide en un jeu de données normalisé."
    )

    doc.add_paragraph("Algorithme d'agrégation :")
    steps = [
        "Chargement des enregistrements bruts depuis csv_historical_data (par paire et timeframe).",
        "Regroupement par (trading_pair_id, date) pour fusionner les doublons.",
        "Calcul : open/close = moyenne pondérée par le volume (ou moyenne simple si volume = 0), high = max, low = min, volume = somme.",
        "Suppression des colonnes intermédiaires (weighted_open, weighted_close).",
        "Insertion en base dans la table cible (ohlcv_minute, ohlcv_hourly ou ohlcv_daily) selon le timeframe.",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ")
        run.bold = True
        p.add_run(step)

    doc.add_heading("3.3 Homogénéisation des formats", level=2)
    doc.add_paragraph(
        "Les dates sont normalisées en datetime Python (UTC). Les colonnes sont standardisées : "
        "date, open, high, low, close, volume_quote, trading_pair_id. "
        "Les entrées corrompues (NaN, valeurs nulles) sont filtrées lors du parsing CSV (pandas). "
        "Les doublons sont gérés par les contraintes d'unicité en base de données."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # PAGE 4 — C4 : BASE DE DONNÉES + RGPD
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("4. C4 — Base de données et conformité RGPD", level=1)

    doc.add_heading("4.1 Choix du SGBD", level=2)
    doc.add_paragraph(
        "PostgreSQL 16 a été retenu pour les raisons suivantes : support natif des contraintes d'unicité "
        "composites, performances sur les requêtes analytiques avec de grands volumes de données temporelles, "
        "compatibilité avec SQLAlchemy ORM, et déploiement simplifié via Docker."
    )

    doc.add_heading("4.2 Modélisation des données", level=2)
    doc.add_paragraph(
        "La modélisation suit la méthode Merise. Le modèle conceptuel (MCD) et le modèle physique (MPD) "
        "sont présentés en Annexe A. Le schéma comporte 12 tables :"
    )
    add_styled_table(doc,
        ["Table", "Description", "Contraintes d'unicité"],
        [
            ["currencies", "Cryptomonnaies et fiats (BTC, ETH, USD…)", "(name, symbol, rank, type)"],
            ["trading_pairs", "Paires de trading (BTCUSDT, ETHUSDT…)", "(base_currency_id, quote_currency_id)"],
            ["exchanges", "Places de marché (Binance)", "(name, slug)"],
            ["cryptocurrency_csv", "Métadonnées des fichiers CSV historiques", "(exchange_id, trading_pair_id, timeframe)"],
            ["csv_historical_data", "Données historiques importées depuis CSV", "(csv_file_id, date)"],
            ["ohlcv_minute", "Données OHLCV à la minute", "(trading_pair_id, date)"],
            ["ohlcv_hourly", "Données OHLCV horaires", "(trading_pair_id, date)"],
            ["ohlcv_daily", "Données OHLCV journalières", "(trading_pair_id, date)"],
            ["predictions_hourly", "Prédictions horaires du modèle ML", "(trading_pair_id, date)"],
            ["predictions_daily", "Prédictions journalières du modèle ML", "(trading_pair_id, date)"],
            ["users", "Utilisateurs de l'API (auth JWT)", "(username) — unique"],
        ],
        col_widths=[4, 7, 5],
    )

    doc.add_heading("4.3 Script d'import et procédure d'installation", level=2)
    doc.add_paragraph(
        "L'import initial est orchestré par init_db_and_data.py, qui enchaîne : création des tables "
        "(Base.metadata.create_all), extraction des données, alimentation de la base, agrégation, "
        "et création de l'utilisateur API. La procédure est reproductible via Docker Compose : "
        "docker compose up db data-api data-scripts -d."
    )

    doc.add_heading("4.4 Conformité RGPD", level=2)
    doc.add_paragraph(
        "Les données de marché (OHLCV, paires de trading) ne constituent pas des données personnelles. "
        "La seule table contenant des données utilisateur est users, qui stocke :"
    )
    rgpd_items = [
        "username : identifiant choisi par l'utilisateur (pas de nom/prénom).",
        "password_hashed : mot de passe hashé avec bcrypt (irréversible).",
        "role et status : métadonnées de gestion.",
    ]
    for item in rgpd_items:
        doc.add_paragraph(f"• {item}")

    doc.add_paragraph(
        "Mesures de conformité mises en œuvre :"
    )
    rgpd_measures = [
        "Minimisation des données : seul le strict nécessaire est stocké (pas d'email, pas de données personnelles).",
        "Hashing irréversible : les mots de passe ne sont jamais stockés en clair (bcrypt).",
        "Droit à l'effacement : l'endpoint DELETE /api/v1/authentification/account permet à l'utilisateur de supprimer son compte.",
        "Registre des traitements : détaillé en Annexe D.",
    ]
    for item in rgpd_measures:
        doc.add_paragraph(f"• {item}")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # PAGE 5 — C5 : API REST
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("5. C5 — API REST et authentification", level=1)

    doc.add_heading("5.1 Architecture de l'API", level=2)
    doc.add_paragraph(
        "L'API est développée avec FastAPI, qui génère automatiquement la documentation OpenAPI (Swagger) "
        "accessible à l'adresse /docs. L'API suit l'architecture REST et expose un CRUD complet."
    )

    doc.add_heading("5.2 Authentification et autorisation", level=2)
    doc.add_paragraph("Le flux d'authentification se déroule ainsi :")
    auth_steps = [
        "L'utilisateur s'inscrit via POST /register (username + password). Le mot de passe est hashé avec bcrypt.",
        "L'utilisateur se connecte via POST /login (OAuth2 Password Flow). Un token JWT est retourné (validité : 30 min).",
        "Les endpoints protégés exigent le header Authorization: Bearer <token>. Le JWT est décodé et l'utilisateur est chargé depuis la base.",
        "Les rôles contrôlent l'accès : user (lecture seule) et script (lecture + écriture/modification/suppression des prédictions).",
        "Un token invalide ou absent retourne une erreur 401. Un rôle insuffisant retourne une erreur 403.",
    ]
    for i, step in enumerate(auth_steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ")
        run.bold = True
        p.add_run(step)

    doc.add_heading("5.3 Endpoints — CRUD complet", level=2)
    doc.add_paragraph(
        "L'API expose 19 endpoints couvrant les 4 opérations CRUD. Le tableau complet est présenté "
        "en Annexe B. Résumé par domaine fonctionnel :"
    )
    add_styled_table(doc,
        ["Domaine", "Create", "Read", "Update", "Delete"],
        [
            ["Utilisateurs", "POST /register", "GET /me", "PUT /password", "DELETE /account"],
            ["Prédictions", "POST /predictions/{freq}", "GET /predictions/…", "PUT /predictions/{freq}/{id}", "DELETE /predictions/{freq}/{id}"],
            ["Données marché", "—", "GET /ohlcv/…, GET /trading_pairs/…", "—", "—"],
            ["Monitoring", "—", "GET /health, GET /metrics", "—", "—"],
        ],
        col_widths=[3, 3.5, 4, 3.5, 3],
    )

    doc.add_heading("5.4 Tests automatisés", level=2)
    doc.add_paragraph(
        "La suite de tests (pytest) comprend 52 tests couvrant l'ensemble des endpoints, "
        "l'authentification, les contrôles de rôles et les fonctions utilitaires. "
        "La couverture des modules API atteint 89 à 100 % selon les fichiers. "
        "Les tests utilisent un mock de la base de données (pas de PostgreSQL requis) et un TestClient FastAPI. "
        "Les résultats sont présentés en Annexe E."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # ANNEXES
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("Annexes", level=1)

    # ── ANNEXE A ──
    doc.add_heading("Annexe A — Modèle Physique des Données (MPD)", level=2)
    doc.add_paragraph(
        "Le modèle physique ci-dessous représente les 12 tables PostgreSQL, leurs colonnes, "
        "types et relations (clés étrangères)."
    )

    mpd_tables = [
        ("currencies", [
            "id : INTEGER (PK)", "name : VARCHAR (NOT NULL)", "symbol : VARCHAR (NOT NULL)",
            "slug : VARCHAR", "sign : VARCHAR", "rank : INTEGER",
            "rank_date : DATETIME", "type : VARCHAR (NOT NULL)",
        ]),
        ("trading_pairs", [
            "id : INTEGER (PK)", "base_currency_id : INTEGER (FK → currencies.id)",
            "quote_currency_id : INTEGER (FK → currencies.id)",
        ]),
        ("exchanges", [
            "id : INTEGER (PK)", "name : VARCHAR (NOT NULL)", "slug : VARCHAR (NOT NULL)",
        ]),
        ("cryptocurrency_csv", [
            "id : INTEGER (PK)", "exchange_id : INTEGER (FK → exchanges.id)",
            "trading_pair_id : INTEGER (FK → trading_pairs.id)",
            "timeframe : VARCHAR (NOT NULL)", "start_date : DATETIME", "end_date : DATETIME",
            "file_url : VARCHAR (NOT NULL)",
        ]),
        ("csv_historical_data", [
            "id : INTEGER (PK)", "csv_file_id : INTEGER (FK → cryptocurrency_csv.id)",
            "date : DATETIME", "open : FLOAT", "high : FLOAT", "low : FLOAT",
            "close : FLOAT", "volume_quote : FLOAT",
        ]),
        ("ohlcv_minute / ohlcv_hourly / ohlcv_daily", [
            "id : INTEGER (PK)", "trading_pair_id : INTEGER (FK → trading_pairs.id)",
            "date : DATETIME", "open : FLOAT", "high : FLOAT", "low : FLOAT",
            "close : FLOAT", "volume_quote : FLOAT",
        ]),
        ("predictions_hourly / predictions_daily", [
            "id : INTEGER (PK)", "trading_pair_id : INTEGER (FK → trading_pairs.id)",
            "date : DATETIME", "predicted_class : INTEGER", "predicted_label : VARCHAR",
            "confidence : FLOAT", "model_name : VARCHAR",
        ]),
        ("users", [
            "id : INTEGER (PK)", "username : VARCHAR (UNIQUE, NOT NULL)",
            "password_hashed : VARCHAR (NOT NULL)", "status : VARCHAR (default: 'active')",
            "role : VARCHAR (default: 'user')",
        ]),
    ]

    for table_name, columns in mpd_tables:
        p = doc.add_paragraph()
        run = p.add_run(table_name)
        run.bold = True
        run.font.size = Pt(10)
        for col in columns:
            p = doc.add_paragraph(f"    {col}")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.name = "Consolas"
        doc.add_paragraph()

    # ── ANNEXE B ──
    doc.add_heading("Annexe B — Tableau complet des endpoints API", level=2)
    add_styled_table(doc,
        ["Méthode", "Endpoint", "Description", "Auth"],
        [
            ["POST", "/api/v1/authentification/login", "Obtenir un token JWT", "Non"],
            ["POST", "/api/v1/authentification/register", "Créer un compte utilisateur", "Non"],
            ["GET", "/api/v1/authentification/me", "Profil de l'utilisateur connecté", "JWT"],
            ["PUT", "/api/v1/authentification/password", "Modifier son mot de passe", "JWT"],
            ["DELETE", "/api/v1/authentification/account", "Supprimer son compte", "JWT"],
            ["GET", "/api/v1/trading_pairs/all", "Paires par symbole de base", "JWT"],
            ["GET", "/api/v1/trading_pairs/trading_pair_by_currency_symbols", "Paire par base/quote", "JWT"],
            ["GET", "/api/v1/ohlcv/minute_by_trading_pair_id", "OHLCV minute", "JWT"],
            ["GET", "/api/v1/ohlcv/hourly_by_trading_pair_id", "OHLCV horaire", "JWT"],
            ["GET", "/api/v1/ohlcv/daily_by_trading_pair_id", "OHLCV journalier", "JWT"],
            ["GET", "/api/v1/predictions/hourly_by_trading_pair_id/{id}", "Prédictions horaires", "JWT"],
            ["GET", "/api/v1/predictions/daily_by_trading_pair_id/{id}", "Prédictions journalières", "JWT"],
            ["GET", "/api/v1/predictions/last_hourly_by_trading_pair_id/{id}", "Dernière prédiction horaire", "JWT"],
            ["GET", "/api/v1/predictions/last_daily_by_trading_pair_id/{id}", "Dernière prédiction journalière", "JWT"],
            ["POST", "/api/v1/predictions/hourly", "Créer une prédiction horaire", "JWT + script"],
            ["POST", "/api/v1/predictions/daily", "Créer une prédiction journalière", "JWT + script"],
            ["PUT", "/api/v1/predictions/hourly/{id}", "Modifier une prédiction horaire", "JWT + script"],
            ["PUT", "/api/v1/predictions/daily/{id}", "Modifier une prédiction journalière", "JWT + script"],
            ["DELETE", "/api/v1/predictions/hourly/{id}", "Supprimer une prédiction horaire", "JWT + script"],
            ["DELETE", "/api/v1/predictions/daily/{id}", "Supprimer une prédiction journalière", "JWT + script"],
            ["GET", "/health", "Healthcheck", "Non"],
            ["GET", "/metrics", "Métriques Prometheus", "Non"],
        ],
        col_widths=[2, 7, 4.5, 2.5],
    )

    # ── ANNEXE C ──
    doc.add_heading("Annexe C — Schéma d'architecture technique", level=2)
    doc.add_paragraph(
        "Le flux de données end-to-end du projet suit le schéma suivant :"
    )
    flow_steps = [
        "Sources externes (Binance API, CoinMarketCap API, CryptoDownload CSV, CoinTelegraph Scrapy)",
        "    ↓",
        "C1_extraction — Scripts d'extraction Python",
        "    ↓",
        "C3_aggregate_ohlcv — Agrégation et nettoyage (pandas, numpy)",
        "    ↓",
        "C4_database — PostgreSQL 16 (SQLAlchemy ORM, 12 tables)",
        "    ↓",
        "C5_api — FastAPI REST API (JWT auth, CRUD, OpenAPI)",
        "    ↓",
        "Consommateurs : modèle ML (Bloc 2), application web (Bloc 3)",
    ]
    for step in flow_steps:
        p = doc.add_paragraph(step)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        for r in p.runs:
            r.font.name = "Consolas"
            r.font.size = Pt(10)

    # ── ANNEXE D ──
    doc.add_heading("Annexe D — Extrait du registre RGPD", level=2)
    add_styled_table(doc,
        ["Traitement", "Données concernées", "Finalité", "Base légale", "Durée de conservation", "Mesures de sécurité"],
        [
            [
                "Gestion des comptes utilisateurs API",
                "username, password_hashed, role, status",
                "Authentification et autorisation d'accès à l'API",
                "Intérêt légitime (sécurisation de l'accès)",
                "Jusqu'à suppression du compte par l'utilisateur (DELETE /account)",
                "Hashing bcrypt, JWT avec expiration 30 min, HTTPS",
            ],
            [
                "Collecte de données de marché",
                "OHLCV, paires de trading, cours",
                "Alimentation du modèle ML de classification",
                "Intérêt légitime",
                "Durée du projet",
                "Données publiques, pas de données personnelles",
            ],
        ],
        col_widths=[3, 3, 3, 2.5, 3, 2.5],
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Procédures de tri : l'endpoint DELETE /api/v1/authentification/account permet à chaque utilisateur "
        "d'exercer son droit à l'effacement. Les données de marché étant publiques et anonymes, "
        "aucune procédure de tri supplémentaire n'est requise."
    )

    # ── ANNEXE E ──
    doc.add_heading("Annexe E — Résultats des tests", level=2)
    doc.add_paragraph("Résultats de l'exécution : 52 tests passés avec succès (0 échec).")
    doc.add_paragraph()

    add_styled_table(doc,
        ["Fichier de test", "Classe", "Nb tests", "Statut"],
        [
            ["test_api_endpoints.py", "TestHealthAndRoot", "2", "PASSED"],
            ["test_api_endpoints.py", "TestTradingPairs", "3", "PASSED"],
            ["test_api_endpoints.py", "TestOHLCV", "4", "PASSED"],
            ["test_api_endpoints.py", "TestPredictions", "4", "PASSED"],
            ["test_api_endpoints.py", "TestPredictionsCRUD", "8", "PASSED"],
            ["test_auth.py", "TestLogin", "5", "PASSED"],
            ["test_auth.py", "TestUserProfile", "6", "PASSED"],
            ["test_auth.py", "TestAuthProtection", "5", "PASSED"],
            ["test_utils.py", "TestValidateDate", "6", "PASSED"],
            ["test_utils.py", "TestParseDate", "3", "PASSED"],
            ["test_utils.py", "TestTimestampConversions", "3", "PASSED"],
            ["test_utils.py", "TestRoundDatetime", "3", "PASSED"],
        ],
        col_widths=[5, 4.5, 2, 2.5],
    )

    doc.add_paragraph()
    doc.add_paragraph("Couverture des modules API :")
    add_styled_table(doc,
        ["Module", "Couverture"],
        [
            ["routes/login.py", "100 %"],
            ["routes/trading_pairs.py", "100 %"],
            ["routes/ohlcv.py", "92 %"],
            ["routes/predictions.py", "89 %"],
            ["utils/auth.py", "95 %"],
            ["utils/classes.py", "100 %"],
            ["utils/functions.py", "100 %"],
        ],
        col_widths=[8, 4],
    )

    # ── SAUVEGARDE ──
    doc.save(str(OUTPUT_PATH))
    print(f"Rapport généré : {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
