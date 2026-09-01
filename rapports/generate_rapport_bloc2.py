"""Génération du rapport professionnel Bloc 2 — Épreuve E2 (C6, C7, C8)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

OUTPUT_PATH = Path(__file__).parent / "Rapport_Bloc2_Veille_Benchmark_Parametrage_IA.docx"


def set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2E74B5")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D9E2F3")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_bullet(doc, title: str, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"• {title} : ")
    run.bold = True
    p.add_run(text)


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        doc.styles[f"Heading {level}"].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # ── PAGE DE GARDE ──
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Rapport professionnel — Accompagnement du choix\n"
        "et intégration d'un service d'intelligence artificielle"
    )
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Titre professionnel Développeur en Intelligence Artificielle — Bloc 2, Épreuve E2\n"
        "Veille, benchmark et paramétrage d'un service de prédiction de tendance crypto"
    )
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Projet support : Classification de Tendance Crypto\n"
        "Certification : RNCP37827 — Développeur en Intelligence Artificielle\n"
        "Compétences visées : C6, C7, C8\n"
        "Dépôt : crypto-certification (GitHub)"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_page_break()

    # ── SOMMAIRE ──
    doc.add_heading("Sommaire", level=1)
    for item in [
        "1. Contexte du projet",
        "2. Organiser une veille technique ciblée (C6)",
        "3. Identifier et recommander un service d'IA à partir du besoin (C7)",
        "4. Paramétrer le service retenu (C8)",
        "5. Conclusion",
        "6. Sources consultées",
        "Annexes",
    ]:
        doc.add_paragraph(item)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 1. CONTEXTE
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("1. Contexte du projet", level=1)

    doc.add_heading("1.1 Le besoin à l'origine du projet", level=2)
    doc.add_paragraph(
        "Le projet « Classification de Tendance Crypto » vise à construire une plateforme d'aide à la "
        "décision pour les investisseurs en cryptomonnaies. L'objectif fonctionnel est de prédire, pour "
        "chaque pas de temps (heure ou jour), si le prix du Bitcoin va monter (UP), descendre (DOWN) "
        "ou rester stable (STABLE), à partir de données de marché historiques et de signaux "
        "complémentaires."
    )
    doc.add_paragraph(
        "La couche données (Bloc 1) collecte et expose via API REST les cours OHLCV, les actualités "
        "Bitcoin scrapées depuis CoinTelegraph, et des indicateurs de sentiment. La question qui ouvre "
        "ce rapport est celle de la couche intelligence artificielle : quel service d'IA préexistant "
        "utiliser pour produire cette prédiction de tendance ?"
    )

    doc.add_heading("1.2 Première hypothèse de travail et ses limites", level=2)
    doc.add_paragraph(
        "En tant qu'alternant ingénieur IA, je manipule quotidiennement des appels API LLM (extraction, "
        "agents, résumé). Ma première hypothèse de travail a donc naturellement été de tester un LLM "
        "généraliste (Anthropic Claude) en zero-shot comme baseline rapide pour la prédiction de "
        "tendance — une démarche que la littérature récente en finance computationnelle utilise elle-même "
        "comme point de comparaison initial avant de proposer des architectures spécialisées (cf. §3.3) — "
        "en lui envoyant les données de marché des 30 derniers jours (Fear & Greed Index + cours BTC en "
        "USD) et en lui demandant de prédire la tendance des prochaines 24 heures. Cette approche est "
        "implémentée dans l'application Streamlit Bloc2_veille/app_benchmark.py."
    )
    doc.add_paragraph("Cette première tentative fonctionne au sens où elle produit une réponse structurée, "
                      "mais révèle quatre limites structurelles :")

    add_styled_table(doc,
        ["Limite", "Constat", "Impact"],
        [
            ["Coût en tokens", "~800–1 200 tokens input par requête (tableau 30 jours)", "~0,003 $/requête — non soutenable à volume élevé"],
            ["Non-reproductibilité", "Réponses variables même à temperature=0", "Impossible de garantir la cohérence des prédictions"],
            ["Pertinence des données", "LLM non entraîné sur données financières récentes", "Précision faible sur une tâche quantitative"],
            ["Empreinte carbone", "Inférence GPU cloud à chaque requête, sans chiffre officiel publié par l'éditeur", "Non quantifiable précisément — argument qualitatif détaillé en §3.5"],
        ],
        col_widths=[3.5, 5.5, 5],
    )

    doc.add_paragraph(
        "Ce constat a motivé une veille formalisée : existe-t-il une solution plus adaptée qu'un LLM "
        "généraliste pour une prédiction quantitative de tendance, tout en conservant le LLM comme "
        "outil complémentaire pour l'analyse qualitative ?"
    )

    doc.add_heading("1.3 Cadrage de l'épreuve E2", level=2)
    doc.add_paragraph(
        "Ce rapport porte sur l'intégration de services d'intelligence artificielle préexistants "
        "(LLM cloud et/ou modèle ML), par opposition au rapport E3 qui couvrira l'exposition du modèle "
        "via API, son intégration dans l'application Django, le monitorage, les tests automatisés et "
        "la chaîne MLOps."
    )
    doc.add_paragraph(
        "Les compétences évaluées ici sont : C6 (veille technique et réglementaire), C7 (benchmark "
        "et recommandation de services IA), C8 (paramétrage du service préconisé)."
    )

    doc.add_heading("1.4 Architecture du projet et périmètre Bloc 2", level=2)
    add_styled_table(doc,
        ["Bloc", "Rôle", "Technologie", "Lien avec ce rapport"],
        [
            ["Bloc1_data", "Collecte et stockage OHLCV + scraping", "FastAPI, PostgreSQL, Scrapy", "Fournit les données d'entrée"],
            ["Bloc2_veille", "Veille + benchmark + POC LLM", "Python, Streamlit, Anthropic SDK", "Périmètre principal E2"],
            ["Bloc3_ml", "Modèle ML retenu (XGBoost)", "FastAPI, scikit-learn, MLflow", "Service préconisé — paramétrage C8"],
            ["Bloc4_app", "Application web utilisateur", "Django", "Hors périmètre E2 (couvert en E3/E4)"],
        ],
        col_widths=[2.5, 4.5, 4, 5],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 2. VEILLE — C6
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("2. Organiser une veille technique ciblée (C6)", level=1)

    doc.add_heading("2.1 Thématique et planification", level=2)
    doc.add_paragraph(
        "La thématique retenue est : « Services d'IA pour la prédiction de tendance crypto — LLM "
        "généraliste vs modèle ML supervisé : coût, tokens, reproductibilité et éco-responsabilité ». "
        "Elle est directement mobilisée dans la mise en situation : il s'agit d'un choix d'outil pour "
        "un besoin concret du projet, pas d'un sujet générique."
    )

    doc.add_paragraph("Le travail de veille s'est déroulé en deux temps :")
    add_styled_table(doc,
        ["Période", "Thématique", "Déclencheur"],
        [
            ["Semaine 1–2, ~1h30 puis ~1h/semaine", "Veille n°1 : indicateurs de sentiment et sources de données marché", "Besoin de données d'entrée pour alimenter une prédiction"],
            ["Semaine 3–4, ~1h/semaine", "Veille n°2 : suivre les nouveautés des services IA (RSS/GitHub) et leur consommation réelle (LiteLLM)", "Suivi manuel via console fournisseur insuffisant : pas de traçabilité automatisée par clé/projet (limite identifiée en §4.2.3)"],
        ],
        col_widths=[4, 6, 6],
    )

    doc.add_paragraph(
        "Le rythme retenu (au minimum une heure hebdomadaire) est documenté dans docs/methodologie_agile.md "
        "(rituels Kanban : planning hebdomadaire, review le vendredi). Ce rythme est cohérent avec la "
        "recommandation du référentiel Simplon pour une veille thématique ciblée."
    )

    doc.add_heading("2.2 Première veille : indicateurs de sentiment et sources de données", level=2)
    doc.add_paragraph(
        "Avant de choisir un service d'IA, il fallait identifier des signaux de marché complémentaires "
        "aux cours OHLCV bruts. La veille a porté sur les indicateurs de sentiment crypto disponibles "
        "via API publique."
    )

    add_styled_table(doc,
        ["Source", "Type", "Données", "Fiabilité"],
        [
            ["Alternative.me", "API REST gratuite", "Fear & Greed Index (score 0–100, classification)", "Éditeur identifié, API stable depuis 2018, données quotidiennes"],
            ["CoinGecko", "API REST gratuite", "Cours BTC/USD historiques 30 jours", "Acteur reconnu, documentation OpenAPI, rate-limit documenté"],
            ["CoinTelegraph", "Scraping Scrapy (Bloc1)", "Actualités Bitcoin (titre, date, catégorie)", "Source média reconnue, scraping respectueux (robots.txt, délai 2s)"],
        ],
        col_widths=[3, 3, 5, 5],
    )

    doc.add_paragraph(
        "Le Fear & Greed Index a été retenu comme indicateur principal de veille car il est directement "
        "mobilisé dans le POC de prédiction (parametrage.py, app_benchmark.py) et dans le feature "
        "engineering du modèle ML (Bloc3). Il mesure le sentiment du marché crypto sur une échelle "
        "0 (peur extrême) à 100 (avidité extrême), publié quotidiennement par Alternative.me."
    )

    doc.add_paragraph(
        "Ces scripts (parametrage.py, app_benchmark.py) sont des outils de collecte de données pour "
        "le produit (Bloc1/C1) — ils alimentent le modèle en signaux de marché, mais ne constituent "
        "pas en soi un outil de veille sur l'écosystème des services d'IA eux-mêmes (cf. §2.3, outils "
        "d'agrégation dédiés à la veille C6)."
    )
    add_bullet(doc, "GitHub Issues / Projects", "Suivi des tâches de veille et partage des synthèses avec le formateur via le board Kanban public.")

    doc.add_heading("2.3 Seconde veille : suivre les services IA, de l'annonce au coût réel", level=2)
    doc.add_paragraph(
        "La comparaison chiffrée détaillée des services (précision, tarifs, contraintes) est traitée "
        "en tant que telle au §3 (benchmark C7) : Anthropic Claude, OpenAI GPT-4o et Mistral Large y "
        "sont comparés à partir de leurs pricing pages officielles. Cette seconde veille répond à une "
        "question complémentaire et opérationnelle, en deux temps : comment rester informé des "
        "nouveautés de ces services une fois le choix fait, et comment suivre ce qu'ils coûtent "
        "réellement à l'usage ?"
    )

    doc.add_paragraph("A. Rester informé des nouveautés (outils d'agrégation) :")
    add_bullet(doc, "Flux RSS des blogs éditeurs (signal direct)", "OpenAI (openai.com/news/rss.xml) et Mistral AI (mistral.ai/news/rss) publient un flux RSS public de leurs annonces modèles/produits — vérifié fonctionnel et intégré dans un onglet dédié de app_benchmark.py.")
    add_bullet(doc, "Flux Atom des releases GitHub (proxy imparfait pour Anthropic)", "Anthropic ne publie aucun flux RSS public sur son site. À défaut, le flux Atom natif des releases du SDK anthropic-sdk-python sert de signal de substitution — une nouvelle version du SDK accompagne souvent, mais pas systématiquement, un nouveau modèle. Ce n'est pas équivalent à une annonce directe, et le rapport le documente comme tel plutôt que de les confondre.")
    add_bullet(doc, "Flux Atom des releases GitHub (signal direct pour XGBoost)", "XGBoost est le modèle/la bibliothèque elle-même : sa release GitHub est un signal direct, sans ambiguïté.")

    doc.add_paragraph(
        "B. Suivre le coût réel à l'usage (déclencheur : limite de monitoring identifiée en §4.2.3 — "
        "le suivi manuel via la console Anthropic ne permet ni traçabilité automatisée par clé/projet, "
        "ni comparaison multi-provider). Cette question fait l'objet d'une pratique professionnelle "
        "personnelle plutôt que d'une découverte en cours de projet : j'utilise LiteLLM depuis un an en "
        "alternance pour ce besoin exact. Sources consultées : documentation officielle "
        "docs.litellm.ai (Spend Tracking, Virtual Keys, Budgets & Rate Limits)."
    )

    doc.add_paragraph(
        "LiteLLM est une bibliothèque et un proxy open-source qui expose une interface unifiée "
        "(format OpenAI) vers plus de 100 fournisseurs de LLM. Son mode proxy (litellm[proxy]) ajoute "
        "des clés API virtuelles avec suivi de dépense et budget par clé, adossés à une base "
        "PostgreSQL, et une UI d'administration. Comparé au suivi manuel (console fournisseur, un "
        "écran par provider, pas de budget configurable par clé), LiteLLM centralise le suivi "
        "token/coût quel que soit le fournisseur derrière — pertinent ici puisque la production réelle "
        "et cette certification n'utilisent pas le même fournisseur (cf. limite identifiée en §4.2.3)."
    )
    doc.add_paragraph(
        "Mise en œuvre : un proxy LiteLLM (image officielle ghcr.io/berriai/litellm) a été déployé en "
        "conteneur Docker, adossé à une base PostgreSQL dédiée (service litellm-proxy, "
        "docker-compose.yml), avec une clé virtuelle scopée au modèle claude-sonnet-4-6 et un budget "
        "maximal configuré. Le POC Streamlit route désormais ses appels au travers de ce proxy "
        "(passthrough Anthropic, base_url configurée dynamiquement) — le détail du paramétrage est "
        "traité en C8, §4.2.5."
    )
    doc.add_paragraph(
        "Fiabilité de la donnée de tarification utilisée pour le calcul du coût : LiteLLM calcule "
        "chaque dépense à partir d'un fichier de référence public et versionné, "
        "model_prices_and_context_window.json (dépôt GitHub BerriAI/litellm), une entrée par couple "
        "modèle/fournisseur. Ce n'est pas une boîte noire : la donnée est vérifiable et recoupable "
        "par n'importe qui. Vérification effectuée pour ce rapport (interrogation directe du fichier "
        "et de l'historique Git le 1er septembre 2026) : dernier commit sur ce fichier daté de la "
        "veille (31 août 2026, mise à jour active) ; l'entrée claude-sonnet-4-6 (provider anthropic) "
        "y indique 3 $ / 15 $ par million de tokens (entrée/sortie) — valeur strictement identique à "
        "celle retenue dans le benchmark du §3, issue indépendamment de la pricing page officielle "
        "Anthropic. Le proxy expose également un endpoint de resynchronisation "
        "(POST /reload/model_cost_map, ou programmée via /schedule/model_cost_map_reload) pour "
        "éviter qu'un redémarrage soit nécessaire à chaque mise à jour tarifaire d'un fournisseur."
    )

    doc.add_heading("2.4 Veille réglementaire", level=2)
    doc.add_paragraph(
        "La veille réglementaire complète la veille technique et couvre deux cadres mobilisés dans le projet :"
    )
    add_bullet(doc, "RGPD", "Registre des traitements documenté (docs/rgpd_registre_traitements.md) — les données OHLCV sont publiques ; seuls les comptes utilisateurs constituent des données personnelles.")
    add_bullet(doc, "MiCA (Markets in Crypto-Assets)", "Régulation UE applicable aux services crypto — pertinente car la plateforme affiche des prédictions de marché ; un disclaimer « aide à la décision, pas conseil financier » a été ajouté au pied de page de l'application (Bloc4_app/templates/base.html).")
    add_bullet(doc, "AI Act (UE)", "Obligations de transparence pour les systèmes d'IA à impact limité — le même disclaimer précise qu'une prédiction est produite par un modèle automatisé, et non par un conseiller humain.")

    doc.add_heading("2.5 Outils et modalités de partage", level=2)
    add_styled_table(doc,
        ["Outil", "Usage", "Justification"],
        [
            ["Script Python (parametrage.py)", "Collecte automatisée Fear & Greed", "Gratuit, reproductible, versionné sur Git"],
            ["Streamlit (app_benchmark.py)", "Visualisation et POC LLM", "Interface accessible, déploiement local sans coût"],
            ["GitHub Projects (Kanban)", "Suivi des tâches de veille", "Traçabilité, partage avec le formateur"],
            ["Markdown (docs/)", "Synthèses formalisées", "Format texte structuré, consultable par lecteur d'écran"],
        ],
        col_widths=[4, 4.5, 7.5],
    )

    doc.add_paragraph(
        "Les synthèses sont rédigées en Markdown (titres hiérarchiques, listes, tableaux en texte), "
        "un format compatible avec les recommandations d'accessibilité Valentin Haüy et Atalan AcceDe "
        "pour les documents numériques."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 3. BENCHMARK — C7
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("3. Identifier et recommander un service d'IA à partir du besoin (C7)", level=1)

    doc.add_heading("3.1 Reformulation du besoin", level=2)
    doc.add_paragraph(
        "Besoin exprimé : disposer d'un outil capable de prédire la tendance du Bitcoin à court terme "
        "(24h) pour alimenter un dashboard d'aide à la décision destiné aux utilisateurs de la plateforme."
    )
    doc.add_paragraph(
        "Reformulation technique : à partir de données de marché structurées (OHLCV historiques, "
        "Fear & Greed Index, actualités), produire une classification ternaire (UP / STABLE / DOWN) "
        "avec un niveau de confiance, en respectant les contraintes suivantes :"
    )

    add_styled_table(doc,
        ["Contrainte", "Détail"],
        [
            ["Latence", "< 5 secondes par prédiction"],
            ["Budget", "Limité — projet étudiant / certification, pas de budget SaaS récurrent"],
            ["Reproductibilité", "Prédictions identiques pour des données identiques"],
            ["Éco-responsabilité", "Minimiser l'empreinte carbone par inférence"],
            ["Explicabilité", "Capacité à justifier la prédiction (feature importance ou raisonnement textuel)"],
        ],
        col_widths=[4, 12],
    )

    doc.add_heading("3.2 Services étudiés et non étudiés", level=2)
    doc.add_paragraph("Services étudiés avec test ou POC effectif :")
    add_bullet(doc, "Anthropic Claude Sonnet 4.6", "POC Streamlit — prédiction à partir de 30 jours de données (Fear & Greed + cours BTC).")
    add_bullet(doc, "XGBoost custom (Bloc3)", "Pipeline ML entraîné sur features OHLCV + indicateurs techniques (RSI, MACD, Bollinger, etc.).")

    doc.add_paragraph("Services étudiés sur documentation/benchmarks, sans test direct :")
    add_bullet(doc, "OpenAI GPT-4o", "Documentation pricing et benchmarks publics LLM vs ML.")
    add_bullet(doc, "Mistral Large", "Documentation pricing et benchmarks de raisonnement tabulaire.")

    doc.add_paragraph("Services non étudiés (raisons explicites) :")
    add_styled_table(doc,
        ["Service écarté", "Raison"],
        [
            ["OpenAI GPT-4o", "Coût prohibitif (~0,01 $/requête), pas de crédits gratuits pour un projet étudiant"],
            ["Mistral Large", "Performances inférieures sur données tabulaires, crédits gratuits limités"],
            ["Services ML SaaS (AWS SageMaker, Google Vertex AI)", "Coût d'infrastructure disproportionné pour un projet solo ; préférence pour un modèle auto-hébergé"],
            ["Random Forest / Logistic Regression (Bloc3)", "Testés comme alternatives internes, inférieurs à XGBoost sur les métriques F1 — conservés comme benchmarks internes"],
        ],
        col_widths=[5, 11],
    )

    doc.add_heading("3.3 Comparaison détaillée — LLM vs ML Custom", level=2)
    add_styled_table(doc,
        ["Critère", "LLM (Claude Sonnet 4.6)", "ML Custom (XGBoost Bloc3)"],
        [
            ["Précision", "Faible — non entraîné sur données financières récentes", "Moyenne à bonne — entraîné sur features OHLCV spécifiques"],
            ["Reproductibilité", "Faible — réponses non déterministes", "Élevée — modèle déterministe à paramètres fixés"],
            ["Latence", "2–5 s (appel réseau + inférence cloud)", "< 0,5 s (inférence CPU locale)"],
            ["Coût par requête", "~0,003 $ (input + output tokens)", "0 $ (auto-hébergé, pas de coût marginal)"],
            ["Éco-responsabilité", "Forte empreinte — datacenter GPU", "Faible — modèle léger, inférence CPU"],
            ["Données d'entraînement", "Corpus généraliste (coupure de connaissance)", "OHLCV crypto récentes, mises à jour par cron"],
            ["Explicabilité", "Raisonnement textuel (non vérifiable)", "Feature importance quantifiable (gain XGBoost)"],
            ["Maintenance", "Aucune — service géré par Anthropic", "Ré-entraînement périodique (pipeline MLOps)"],
        ],
        col_widths=[3.5, 6.25, 6.25],
    )

    doc.add_paragraph(
        "La ligne « Précision » s'appuie sur la littérature académique récente consacrée à "
        "l'usage de LLM comme outil de prédiction financière, et pas seulement sur le constat "
        "empirique du POC : Crisostomo & Mykhalyuk (2026) montrent que des LLM généralistes "
        "utilisés sans spécialisation ni supervision humaine forte souffrent d'erreurs de "
        "raisonnement récurrentes sur des tâches d'investissement ; Wang et al. (2024, StockTime) "
        "constatent que les LLM financiers non spécialisés « négligent les caractéristiques "
        "essentielles des séries temporelles », d'où la nécessité d'architectures dédiées pour "
        "obtenir une précision satisfaisante. Ces deux constats convergent avec le choix retenu "
        "ici : XGBoost, entraîné spécifiquement sur les séries temporelles OHLCV, plutôt qu'un LLM "
        "généraliste utilisé tel quel."
    )

    doc.add_heading("3.4 Adéquation par ensemble fonctionnel", level=2)
    add_styled_table(doc,
        ["Fonction", "LLM (Claude)", "ML Custom (XGBoost)"],
        [
            ["Prédiction quantitative de tendance", "Non adapté", "Conçu pour — couvre le besoin principal"],
            ["Analyse qualitative de marché", "Pertinent", "Non prévu"],
            ["Synthèse d'actualités", "Pertinent", "Non prévu"],
            ["Reproductibilité des résultats", "Non garanti", "Garanti"],
            ["Explication du raisonnement", "Textuel, non vérifiable", "Feature importance quantifiable"],
        ],
        col_widths=[5, 5.5, 5.5],
    )

    doc.add_heading("3.5 Démarche éco-responsable", level=2)
    doc.add_paragraph(
        "Aucun des services comparés (Anthropic, OpenAI, Mistral) ne publie de chiffre officiel de "
        "consommation énergétique ou d'empreinte carbone par requête — c'est une limite générale du "
        "secteur, pas propre à ce projet, et je ne dispose donc pas d'un chiffre sourcé à comparer. "
        "Les ordres de grandeur qui circulent dans la littérature généraliste sur l'inférence LLM "
        "cloud vs. un modèle CPU local n'ont pas été retenus ici faute d'avoir pu les recouper avec "
        "au moins deux sources indépendantes fiables."
    )
    doc.add_paragraph(
        "Le raisonnement qualitatif que je peux tenir, sur la base des informations disponibles, est "
        "le suivant : l'architecture retenue (XGBoost local pour l'essentiel des prédictions, LLM "
        "cloud restreint à un usage complémentaire ponctuel) sollicite structurellement moins de "
        "calcul distant qu'une architecture « tout LLM » où chaque prédiction interrogerait un modèle "
        "de fondation multimodal : moins d'appels réseau, une inférence CPU légère (quelques "
        "millisecondes pour XGBoost) plutôt qu'une inférence GPU cloud à chaque requête, et un "
        "ré-entraînement du modèle ML limité à quelques minutes CPU par cycle. C'est un argument de "
        "sobriété par conception, pas un résultat mesuré."
    )

    doc.add_heading("3.6 Contraintes techniques et pré-requis", level=2)
    doc.add_paragraph("LLM (Anthropic Claude) :")
    add_bullet(doc, "Clé API", "ANTHROPIC_API_KEY — jamais commitée, lue depuis .env")
    add_bullet(doc, "Dépendances", "SDK Python anthropic, streamlit, pandas, requests")
    add_bullet(doc, "Connectivité", "HTTPS sortant obligatoire vers api.anthropic.com")
    add_bullet(doc, "Rate limit", "1 000 req/min sur le tier gratuit")

    doc.add_paragraph("ML Custom (XGBoost Bloc3) :")
    add_bullet(doc, "Infrastructure", "Docker Compose (pipeline + API + MLflow)")
    add_bullet(doc, "Données", "Pipeline Bloc1 fonctionnel (API OHLCV alimentant les features)")
    add_bullet(doc, "Ré-entraînement", "Cron horaire et journalier (entrypoint.sh)")
    add_bullet(doc, "Dépendances", "scikit-learn, xgboost, pandas, pandas-ta, MLflow")

    doc.add_heading("3.7 Conclusions du benchmark", level=2)
    doc.add_paragraph(
        "Résultats réels du backtest XGBoost (granularité journalière, période de test "
        "2025-01-01 → 2025-05-31, métriques trackées dans MLflow après exécution du pipeline) :"
    )
    add_styled_table(doc,
        ["Paire", "Accuracy", "F1 macro", "Direction accuracy"],
        [
            ["BTC-USD", "0,4133", "0,3778", "0,3889"],
            ["BTC-USDT", "0,4533", "0,4292", "0,4184"],
        ],
        col_widths=[4, 4, 4, 4],
    )
    doc.add_paragraph(
        "Ces résultats sont modestes mais réels et reproductibles : une accuracy de 41 à 45 % sur "
        "un problème à 3 classes (le hasard pur donnerait environ 33 % sur des classes équilibrées), "
        "cohérente avec la difficulté connue de prédire un mouvement de prix crypto à courte échéance. "
        "Le F1 par classe (disponible dans MLflow) montre en particulier que le modèle est plus fiable "
        "sur la détection de baisse/hausse marquée que sur la classe STABLE — un point identifié pour "
        "une itération future (ajustement du seuil de classification ou rééquilibrage des classes)."
    )
    doc.add_paragraph(
        "Le modèle XGBoost est évalué avec un protocole de backtest formel — fenêtres de test "
        "glissantes, métriques accuracy/f1_macro/direction_accuracy trackées dans MLflow sur la "
        "période 2025-01-01 à 2025-05-31. Le LLM, lui, n'a volontairement pas été soumis au même "
        "protocole de backtest systématique : chaque prédiction coûte un appel API payant, et la "
        "littérature récente (Crisostomo & Mykhalyuk 2026 ; Wang et al. 2024) montre que des LLM "
        "généralistes en zero-shot obtiennent une précision proche du hasard sur ce type de tâche — "
        "un backtest formel aurait probablement confirmé ce point sans justifier son coût. C'est "
        "précisément cette asymétrie de rigueur, mesurable pour l'un et coûteuse à mesurer pour "
        "l'autre sans gain attendu, qui motive le choix de XGBoost comme service de production."
    )
    doc.add_paragraph(
        "Service retenu pour la prédiction de tendance : XGBoost (modèle custom Bloc3). Il est "
        "déterministe, reproductible, sans coût marginal, et sollicite structurellement moins de "
        "calcul distant qu'un appel LLM cloud (cf. §3.5 pour l'argument qualitatif détaillé — aucun "
        "chiffre officiel d'empreinte carbone n'est disponible pour un ratio précis)."
    )
    doc.add_paragraph(
        "Le LLM (Anthropic Claude) reste pertinent comme outil complémentaire pour : la synthèse "
        "qualitative de conditions de marché, l'explication textuelle des tendances à destination "
        "d'utilisateurs non techniques, et l'enrichissement de l'interface (POC Streamlit). "
        "Aucun des deux services pris isolément ne couvre le besoin complet : l'un ne garantit pas "
        "la reproductibilité quantitative, l'autre ne produit pas d'analyse en langage naturel."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 4. PARAMÉTRAGE — C8
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("4. Paramétrer le service retenu (C8)", level=1)

    doc.add_paragraph(
        "Le paramétrage couvre les deux composants de l'architecture retenue : le modèle ML XGBoost "
        "(service principal) et le LLM Claude (service complémentaire, POC démontrable)."
    )

    doc.add_heading("4.1 Paramétrage du service principal — XGBoost (Bloc3)", level=2)

    doc.add_heading("4.1.1 Environnement d'exécution", level=3)
    add_styled_table(doc,
        ["Composant", "Configuration"],
        [
            ["Image Docker pipeline", "pipeline.Dockerfile — Python 3.11, cron horaire + journalier"],
            ["Image Docker API ML", "api.Dockerfile — FastAPI port 8002"],
            ["Image Docker MLflow", "mlflow.Dockerfile — UI monitoring port 5000"],
            ["Orchestration", "docker-compose.yml — volumes partagés ml_models, mlruns"],
        ],
        col_widths=[4.5, 11.5],
    )

    doc.add_heading("4.1.2 Configuration du modèle", level=3)
    doc.add_paragraph(
        "Les hyperparamètres et la sélection d'algorithme sont externalisés dans des fichiers YAML "
        "versionnés sur Git :"
    )
    add_bullet(doc, "config/hour_models_config.yaml", "Modèle horaire par paire de trading (XGBClassifier par défaut).")
    add_bullet(doc, "config/day_models_config.yaml", "Modèle journalier par paire de trading.")
    add_bullet(doc, "config/ml_config.yaml", "Seuil de classification (0,5 %), fenêtres d'entraînement/test.")
    add_bullet(doc, "config/data_config.yaml", "Paramètres de récupération des OHLCV depuis l'API Bloc1.")

    add_styled_table(doc,
        ["Paramètre", "Valeur", "Justification"],
        [
            ["Algorithme", "XGBClassifier", "Meilleures métriques F1 sur données tabulaires crypto"],
            ["Classes", "0=DOWN, 1=STABLE, 2=UP", "Classification ternaire avec seuil 0,5 %"],
            ["Features", "RSI, MACD, Bollinger, SMA, EMA, lags, rendements", "Indicateurs techniques standard en analyse crypto"],
            ["Période entraînement BTCUSDT daily", "2020-01-01 → 2024-12-31", "5 ans de données pour robustesse"],
            ["Période test", "2025-01-01 → 2025-05-31", "Validation out-of-sample"],
        ],
        col_widths=[4.5, 4.5, 7],
    )

    doc.add_heading("4.1.3 Pipeline et ré-entraînement", level=3)
    doc.add_paragraph(
        "Le script update_models_and_predictions.py exécute le pipeline complet : récupération OHLCV "
        "via JWT Bloc1 → feature engineering → entraînement → évaluation → prédiction → envoi vers "
        "Bloc1 → sauvegarde pickle. Il est déclenché par cron (entrypoint.sh) : horaire à XX:05, "
        "journalier à 00:10."
    )

    doc.add_heading("4.1.4 Monitorage (MLflow) et accès", level=3)
    doc.add_paragraph(
        "MLflow (port 5000) track les métriques accuracy, f1_macro et direction_accuracy à chaque "
        "exécution du pipeline. C'est le monitorage natif du service ML retenu — accessible via "
        "http://localhost:5000 après docker compose up mlflow-server."
    )
    doc.add_paragraph(
        "Côté accès : l'API de classification (ml-api) est protégée par authentification JWT "
        "(Bloc3_ml/src/api/routes/classify.py::get_current_user, Bloc3_ml/src/api/utils/deps.py), "
        "cohérente avec le système d'authentification de l'API Bloc1. L'interface MLflow, elle, "
        "n'est aujourd'hui pas authentifiée (mlflow.Dockerfile ne configure aucun mécanisme d'accès) "
        "— c'est un point de vigilance identifié, acceptable en environnement de développement local "
        "mais à corriger avant toute exposition au-delà de ce cadre (par exemple via un reverse-proxy "
        "avec authentification, ou la variable MLFLOW_TRACKING_TOKEN)."
    )

    doc.add_heading("4.2 Paramétrage du service complémentaire — Claude (POC Streamlit)", level=2)

    doc.add_heading("4.2.1 Configuration de l'appel API", level=3)
    add_styled_table(doc,
        ["Paramètre", "Valeur", "Justification"],
        [
            ["Modèle", "claude-sonnet-4-6", "Modèle Sonnet actuel (remplace claude-sonnet-4-20250514 déprécié)"],
            ["Temperature", "0.0", "Réponses les plus déterministes possibles"],
            ["Max tokens", "256", "Réponse courte structurée (prédiction + confiance + raison)"],
            ["Format de réponse", "PREDICTION / CONFIANCE / RAISON", "Parsing automatique dans app_benchmark.py"],
        ],
        col_widths=[3, 4, 9],
    )

    doc.add_heading("4.2.2 Construction du prompt", level=3)
    doc.add_paragraph(
        "Le prompt envoie au LLM un tableau fusionné (Fear & Greed + cours BTC) sur 30 jours, "
        "trié du plus récent au plus ancien. Il contraint le modèle à répondre strictement au format "
        "PREDICTION / CONFIANCE / RAISON, parsé par la fonction _parse_llm_response()."
    )
    doc.add_paragraph(
        "Exemple de prompt (extrait) :\n"
        "« Tu es un analyste crypto. Voici les données des 30 derniers jours pour Bitcoin : "
        "Fear & Greed Index + Cours BTC en USD. [tableau]. En te basant uniquement sur ces données, "
        "prédis la tendance du Bitcoin pour les prochaines 24h. Réponds STRICTEMENT au format : "
        "PREDICTION: [HAUSSE ou BAISSE ou STABLE] / CONFIANCE: [0-100] / RAISON: [phrase]. »"
    )

    doc.add_heading("4.2.3 Mesures collectées à chaque appel", level=3)
    add_styled_table(doc,
        ["Métrique", "Usage"],
        [
            ["input_tokens / output_tokens", "Estimation du coût par requête (tarif Anthropic Sonnet 4.6)"],
            ["latency_s", "Comparaison avec latence ML local (< 0,5 s)"],
            ["prediction / confidence / reason", "Résultat structuré affiché dans Streamlit"],
        ],
        col_widths=[5, 11],
    )
    doc.add_paragraph(
        "Ce suivi applicatif est complété par deux niveaux de monitorage externes : la console "
        "Anthropic (console.anthropic.com), monitorage natif du service SaaS, consultable "
        "manuellement par fournisseur ; et le proxy LiteLLM (§4.2.5), qui automatise ce suivi par "
        "clé API et centralise plusieurs fournisseurs — c'est l'outil effectivement utilisé pour "
        "vérifier que le coût réel par requête reste conforme à l'estimation faite dans le "
        "benchmark (C7)."
    )

    doc.add_heading("4.2.4 Gestion des accès et sécurité", level=3)
    add_bullet(doc, "Clé API", "ANTHROPIC_API_KEY lue depuis .env via python-dotenv — jamais commitée (.gitignore).")
    add_bullet(doc, "Saisie fallback", "Si la clé n'est pas en .env, Streamlit propose un champ password pour saisie manuelle.")
    add_bullet(doc, "Gestion d'erreurs", "anthropic.APIError capturée et affichée dans l'interface sans crash de l'application.")

    doc.add_heading("4.2.5 Paramétrage du proxy LiteLLM (monitoring token/coût, veille C6)", level=3)
    doc.add_paragraph(
        "Suite à la veille du §2.3, un proxy LiteLLM a été déployé et paramétré comme service "
        "Docker du projet (litellm-proxy, docker-compose.yml), plutôt que documenté comme simple "
        "piste théorique :"
    )
    add_bullet(doc, "Image", "ghcr.io/berriai/litellm:v1.90.2 (version épinglée), backend Prisma/PostgreSQL dédié (base litellm sur l'instance Postgres du projet).")
    add_bullet(doc, "Configuration du modèle", "litellm_config.yaml déclare claude-sonnet-4-6 (anthropic/claude-sonnet-4-6, clé lue depuis ANTHROPIC_API_KEY) — un seul modèle pour ce POC, extensible aux autres fournisseurs du benchmark sans changer le code applicatif.")
    add_bullet(doc, "Authentification", "LITELLM_MASTER_KEY (accès admin du proxy) et LITELLM_SALT_KEY (chiffrement des identifiants stockés), générées aléatoirement et lues depuis .env — jamais commitées.")
    add_bullet(doc, "Clé virtuelle applicative", "Générée via POST /key/generate (scope : modèle claude-sonnet-4-6, budget maximal configuré) — c'est cette clé, et non la clé Anthropic brute, qui est utilisée par app_benchmark.py.")
    add_bullet(doc, "Intégration applicative", "call_anthropic() route désormais l'appel via le SDK anthropic standard, avec base_url pointant vers le proxy (endpoint passthrough /anthropic) — aucune dépendance supplémentaire, repli automatique sur l'appel direct si le proxy n'est pas configuré.")

    doc.add_paragraph(
        "Vérification en conditions réelles : un appel passant par le proxy a été effectué et "
        "tracé — dépense de 0,000117 $ pour 14 tokens d'entrée / 5 de sortie, associée à la clé "
        "virtuelle, consultable via GET /key/info et GET /spend/logs. Un second appel confirme "
        "l'incrémentation correcte du suivi (dépense cumulée passée à 0,000561 $), avec un délai "
        "d'agrégation asynchrone de quelques secondes propre au proxy. Ce suivi est également "
        "affiché dans l'application (panneau « Observabilité — veille C6 » de app_benchmark.py) et "
        "consultable via l'UI d'administration du proxy (/ui, authentifiée par la clé maître)."
    )
    doc.add_paragraph(
        "Point de vigilance : le déploiement initial du proxy s'est heurté à une contrainte "
        "mémoire de l'environnement Docker local (limite de VM insuffisante pour la génération du "
        "client Prisma pendant l'initialisation, provoquant un OOM-kill silencieux) — corrigé en "
        "libérant de la mémoire sur l'hôte plutôt qu'en modifiant le paramétrage du service. Ce "
        "type d'incident de déploiement, comme celui documenté pour Bloc3/MLflow, est documenté "
        "tel quel plutôt que masqué."
    )

    doc.add_heading("4.3 Données d'entrée communes", level=2)
    add_styled_table(doc,
        ["Source", "Module", "Fréquence de mise à jour"],
        [
            ["Fear & Greed Index", "parametrage.py → API Alternative.me", "Temps réel à chaque chargement Streamlit"],
            ["Cours BTC 30j", "app_benchmark.py → API CoinGecko", "Temps réel à chaque chargement"],
            ["Actualités BTC", "Bloc1 scraping → scraped_articles.json", "Manuel (scrapy crawl) ou bouton « Rafraîchir » dans Streamlit"],
            ["OHLCV historiques", "Bloc1 API → Bloc3 pipeline", "Cron horaire + journalier"],
        ],
        col_widths=[3.5, 5.5, 7],
    )

    doc.add_heading("4.4 Documentation et procédures", level=2)
    add_styled_table(doc,
        ["Document", "Contenu"],
        [
            ["Bloc2_veille/README.md", "Installation (uv sync), lancement veille et benchmark Streamlit"],
            ["Bloc3_ml/README.md", "Architecture pipeline, config YAML, MLflow, API classification"],
            ["docs/benchmark_services_ia.md", "Benchmark formel C7 — expression de besoin, comparaison, conclusions"],
            ["docs/methodologie_agile.md", "Planification veille (rituels hebdomadaires)"],
            ["litellm_config.yaml", "Configuration du proxy LiteLLM (modèle, clé API source)"],
            [".env.example", "Variables ANTHROPIC_API_KEY, LITELLM_MASTER_KEY/SALT_KEY/PROXY_URL/VIRTUAL_KEY documentées"],
        ],
        col_widths=[5, 11],
    )

    doc.add_paragraph(
        "La documentation est rédigée en Markdown (titres hiérarchiques, listes, blocs de code), "
        "format consultable par un lecteur d'écran sans dépendre d'une mise en forme visuelle."
    )

    doc.add_heading("4.5 Procédures d'installation et de test", level=2)
    doc.add_paragraph("Veille Fear & Greed :")
    doc.add_paragraph("cd Bloc2_veille && uv sync && uv run python parametrage.py", style="Intense Quote")

    doc.add_paragraph("POC LLM Streamlit :")
    doc.add_paragraph(
        "export ANTHROPIC_API_KEY=sk-ant-...\n"
        "cd Bloc2_veille && uv run streamlit run app_benchmark.py",
        style="Intense Quote",
    )

    doc.add_paragraph("Pipeline ML (service retenu) :")
    doc.add_paragraph("docker compose up ml-pipeline mlflow-server ml-api -d", style="Intense Quote")

    doc.add_paragraph("Proxy LiteLLM (veille C6, monitoring token/coût) :")
    doc.add_paragraph(
        "docker compose up litellm-proxy -d\n"
        "curl -X POST http://localhost:4010/key/generate \\\n"
        "  -H \"Authorization: Bearer $LITELLM_MASTER_KEY\" \\\n"
        "  -d '{\"models\": [\"claude-sonnet-4-6\"], \"max_budget\": 5}'",
        style="Intense Quote",
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 5. CONCLUSION
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("5. Conclusion", level=1)
    doc.add_paragraph(
        "Ce projet illustre un chemin d'intégration d'un service d'IA préexistant qui n'a pas été "
        "linéaire : un premier essai naïf (envoyer les données de marché à un LLM cloud) a fonctionné "
        "assez pour révéler ses propres limites — coût en tokens, non-reproductibilité, faible pertinence "
        "sur une tâche quantitative — ce qui a orienté une veille formalisée vers un modèle ML supervisé "
        "auto-hébergé, tout en conservant le LLM comme outil complémentaire d'analyse qualitative."
    )
    doc.add_paragraph(
        "Le service finalement paramétré n'est donc pas le premier candidat testé, mais le résultat "
        "de deux itérations de remise en question, chacune documentée par une veille distincte et "
        "déclenchée par un constat concret plutôt que par une préférence a priori."
    )
    doc.add_paragraph(
        "L'architecture retenue — XGBoost local pour la prédiction + Claude en POC complémentaire — "
        "répond aux contraintes de coût, reproductibilité et éco-responsabilité identifiées dans le "
        "benchmark, tout en laissant une porte ouverte à l'enrichissement de l'interface utilisateur "
        "par des synthèses en langage naturel."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 6. SOURCES
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("6. Sources consultées", level=1)
    sources = [
        "Alternative.me — Fear & Greed Index API (documentation, source primaire)",
        "Anthropic — Claude Platform Pricing et Model Deprecations (source primaire, éditeur)",
        "OpenAI — API Pricing GPT-4o (source primaire, éditeur)",
        "Mistral AI — Documentation et tarification (source primaire, éditeur)",
        "CoinGecko — API v3 market_chart (documentation, source primaire)",
        "LiteLLM — Spend Tracking, Virtual Keys, Simple Proxy / AI Gateway (docs.litellm.ai, source primaire, éditeur)",
        "Règlement spécifique RNCP37827 — Développeur en Intelligence Artificielle (Simplon, 2023)",
        "Référentiel Activités Compétences et évaluation — Bloc 2 (Simplon, 2023)",
        "Règlement UE MiCA — Markets in Crypto-Assets (cadre réglementaire crypto UE)",
        "Règlement UE AI Act — Intelligence artificielle (obligations transparence)",
        "Commission Nationale Informatique et Libertés (CNIL) — Guide RGPD",
        "Valentin Haüy — Recommandations accessibilité documents numériques",
        "Atalan AcceDe — Guide accessibilité des contenus web",
        "XGBoost Documentation — Parameters et Python API (source primaire)",
        "MLflow Documentation — Tracking et Model Registry (source primaire)",
        "Crisostomo, R. & Mykhalyuk, D. (2026) — Large Language Models and Stock Investing: Is the Human Factor Required? (arXiv:2603.19944, vérifié)",
        "Wang, S. et al. (2024) — StockTime: A Time Series Specialized Large Language Model Architecture for Stock Price Prediction (arXiv:2409.08281, vérifié)",
        "A Review of Large Language Models for Stock Price Forecasting from a Hedge-Fund Perspective (arXiv:2605.05211)",
        "Empowering Time Series Analysis with Large Language Models: A Survey (arXiv:2402.03182)",
        "Sources internes : docs/benchmark_services_ia.md, Bloc2_veille/, Bloc3_ml/, docs/methodologie_agile.md",
    ]
    for s in sources:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # ANNEXES
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading("Annexes", level=1)

    doc.add_heading("Annexe A — Expression de besoin (extrait benchmark)", level=2)
    doc.add_paragraph(
        "Objectif : prédire la tendance du Bitcoin (hausse / baisse / stable) à court terme (24h) "
        "pour alimenter un outil d'aide à la décision.\n"
        "Contraintes : données OHLCV + Fear & Greed Index ; latence < 5 s ; budget limité ; "
        "reproductibilité ; éco-responsabilité."
    )

    doc.add_heading("Annexe B — Grille d'adéquation fonctionnelle", level=2)
    add_styled_table(doc,
        ["Fonction", "LLM", "ML Custom"],
        [
            ["Prédiction quantitative", "Non adapté", "Conçu pour"],
            ["Analyse qualitative", "Pertinent", "Non prévu"],
            ["Synthèse actualités", "Pertinent", "Non prévu"],
            ["Reproductibilité", "Non garanti", "Garanti"],
            ["Explicabilité", "Textuel", "Feature importance"],
        ],
        col_widths=[4.5, 5.75, 5.75],
    )

    doc.add_heading("Annexe C — Exemple de prompt LLM", level=2)
    doc.add_paragraph(
        "Tu es un analyste crypto. Voici les données des 30 derniers jours pour Bitcoin :\n"
        "- Fear & Greed Index (score de 0 à 100)\n"
        "- Cours BTC en USD\n\n"
        "[tableau fusionné date | score | sentiment | price]\n\n"
        "En te basant uniquement sur ces données, prédis la tendance du Bitcoin pour les prochaines 24h.\n"
        "Réponds STRICTEMENT au format :\n"
        "PREDICTION: [HAUSSE ou BAISSE ou STABLE]\n"
        "CONFIANCE: [pourcentage entre 0 et 100]\n"
        "RAISON: [une phrase courte]"
    )

    doc.add_heading("Annexe D — Configuration YAML Bloc3 (extrait)", level=2)
    doc.add_paragraph(
        "Fichier config/ml_config.yaml :\n"
        "  classification_threshold: 0.005  # 0,5 % de variation\n"
        "  classes: {down: 0, stable: 1, up: 2}\n\n"
        "Fichier config/day_models_config.yaml :\n"
        "  BTCUSDT:\n"
        "    model: XGBClassifier\n"
        "    train_start: '2020-01-01'\n"
        "    train_end: '2024-12-31'\n"
        "    test_start: '2025-01-01'\n"
        "    test_end: '2025-05-31'"
    )

    doc.add_heading("Annexe E — Commandes de lancement", level=2)
    add_styled_table(doc,
        ["Action", "Commande"],
        [
            ["Veille Fear & Greed", "cd Bloc2_veille && uv run python parametrage.py"],
            ["POC LLM Streamlit", "cd Bloc2_veille && uv run streamlit run app_benchmark.py"],
            ["Scraping actualités", "cd Bloc1_data && uv run scrapy crawl cointelegraph"],
            ["Pipeline ML + MLflow", "docker compose up ml-pipeline mlflow-server ml-api -d"],
            ["Régénérer ce rapport", "cd rapports && .venv/bin/python generate_rapport_bloc2.py"],
        ],
        col_widths=[4, 12],
    )

    doc.add_heading("Annexe F — Cahier des charges du projet (extraits)", level=2)
    doc.add_paragraph(
        "Transcription des extraits pertinents du cahier des charges du projet "
        "(docs/cahier_des_charges_classification_tendance.md), à l'origine du besoin fonctionnel "
        "reformulé en §3.1."
    )
    doc.add_paragraph(
        "1.1 Objectif : Construire une plateforme complète de classification de tendance pour les "
        "cryptomonnaies. Le système prédit, pour chaque pas de temps (heure ou jour), si le prix va "
        "monter (UP), descendre (DOWN) ou rester stable (STABLE). C'est un problème de classification "
        "supervisée à 3 classes, contrairement à un problème de régression qui prédirait le prix exact."
    )
    doc.add_paragraph(
        "1.2 Paires de trading suivies : BTC-USD (Bitcoin / US Dollar) et BTC-USDT (Bitcoin / Tether "
        "USDt), granularités horaire et journalière."
    )
    doc.add_paragraph(
        "1.4 Définition de la variable cible : la cible est construite à partir de la variation "
        "relative du prix de clôture entre deux pas de temps consécutifs "
        "(variation = (close[t] - close[t-1]) / close[t-1]), avec un seuil configurable (0,5 % par "
        "défaut) : UP si variation > +seuil, DOWN si variation < -seuil, STABLE sinon. Le seuil est "
        "paramétrable en YAML pour être ajusté sans modifier le code."
    )

    doc.save(str(OUTPUT_PATH))
    print(f"Rapport généré : {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
